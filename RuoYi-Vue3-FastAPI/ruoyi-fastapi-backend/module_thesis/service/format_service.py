"""
论文格式化服务 - 使用AI读取Word文档并生成格式化指令，然后进行格式化
"""
import json
import os
import tempfile
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, TYPE_CHECKING
from sqlalchemy.ext.asyncio import AsyncSession

from exceptions.exception import ServiceException
from module_thesis.dao.template_dao import UniversalInstructionSystemDao
from module_thesis.service.ai_generation_service import AiGenerationService
from utils.log_util import logger

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
    # 尝试获取版本信息
    try:
        import docx
        docx_version = getattr(docx, '__version__', '未知版本')
        logger.info(f"python-docx 已成功导入，版本: {docx_version}")
    except:
        logger.info("python-docx 已成功导入（无法获取版本信息）")
except ImportError as e:
    DOCX_AVAILABLE = False
    # 为了类型注解，定义一个占位符类型
    if TYPE_CHECKING:
        from docx import Document
    else:
        Document = None  # type: ignore
    import sys
    import os
    python_path = sys.executable
    python_version = sys.version
    error_detail = str(e)
    logger.error(f"python-docx 导入失败！")
    logger.error(f"  Python路径: {python_path}")
    logger.error(f"  Python版本: {python_version}")
    logger.error(f"  错误详情: {error_detail}")
    logger.error(f"  请检查是否在正确的虚拟环境中运行，并执行: pip install python-docx")
    logger.warning(f"python-docx 未安装，Word文档处理功能将不可用。请运行: pip install python-docx。错误详情: {error_detail}")


class FormatService:
    """
    论文格式化服务类
    """
    
    @classmethod
    async def read_word_document_with_ai(
        cls,
        query_db: AsyncSession,
        word_file_path: str,
        config_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        使用AI读取Word文档并提取格式指令
        
        :param query_db: 数据库会话
        :param word_file_path: Word文档路径
        :param config_id: AI模型配置ID（可选）
        :return: 格式指令和文档分析结果
        """
        if not DOCX_AVAILABLE:
            import sys
            python_path = sys.executable
            error_msg = (
                f'python-docx 未安装，无法处理Word文档。\n'
                f'当前Python路径: {python_path}\n'
                f'请执行以下命令安装: pip install python-docx\n'
                f'如果使用虚拟环境，请确保已激活虚拟环境。'
            )
            logger.error(error_msg)
            raise ServiceException(message=error_msg)
        
        if not os.path.exists(word_file_path):
            raise ServiceException(message=f'Word文档不存在: {word_file_path}')
        
        # 检查文件格式：python-docx 只能处理 .docx 格式，不能处理 .doc 格式
        file_ext = os.path.splitext(word_file_path)[1].lower()
        if file_ext == '.doc':
            raise ServiceException(
                message=f'模板文件是 .doc 格式，无法直接处理。python-docx 只能处理 .docx 格式。'
                f'请将模板文件转换为 .docx 格式后重新上传。文件路径: {word_file_path}'
            )
        elif file_ext != '.docx':
            raise ServiceException(
                message=f'模板文件格式不支持（当前格式: {file_ext}）。请使用 .docx 格式的Word文档。文件路径: {word_file_path}'
            )
        
        try:
            print(f"[读取Word文档] 开始处理文件: {word_file_path}")
            print(f"  开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            logger.info(f"[读取Word文档] 开始处理文件: {word_file_path}")
            logger.info(f"  开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # 读取Word文档并提取文本内容
            print("=" * 100)
            print(f"[读取Word文档] 步骤1/2: 打开Word文档并提取文本内容...")
            print(f"  开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            print(f"  文件路径: {word_file_path}")
            logger.info("=" * 100)
            logger.info(f"[读取Word文档] 步骤1/2: 打开Word文档并提取文本内容...")
            logger.info(f"  开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            logger.info(f"  文件路径: {word_file_path}")
            
            print(f"  正在打开Word文档...")
            logger.info(f"  正在打开Word文档...")
            doc = Document(word_file_path)
            para_count = len(doc.paragraphs)
            print(f"  ✓ Word文档打开成功，段落数: {para_count}")
            logger.info(f"  ✓ Word文档打开成功，段落数: {para_count}")
            
            print(f"  正在提取文档文本内容（保留所有段落，包括空行）...")
            logger.info(f"  正在提取文档文本内容（保留所有段落，包括空行）...")
            # 提取文档文本内容（保留所有段落，包括空行）
            document_text = cls._extract_document_text(doc)
            text_length = len(document_text)
            print(f"  ✓ 文本提取完成")
            print(f"  段落数: {para_count}")
            print(f"  文本长度: {text_length} 字符 ({text_length / 1024:.2f} KB)")
            print(f"  完成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            logger.info(f"  ✓ 文本提取完成")
            logger.info(f"  段落数: {para_count}")
            logger.info(f"  文本长度: {text_length} 字符 ({text_length / 1024:.2f} KB)")
            logger.info(f"  完成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # 输出提取的文档文本内容（用于调试，显示前1000字符）
            preview_length = min(1000, text_length)
            text_preview = document_text[:preview_length]
            if text_length > preview_length:
                text_preview += f"\n\n[文档内容过长，已截断前{preview_length}字符，总长度: {text_length}字符]"
            print("=" * 100)
            print("[读取Word文档] 📄 提取的文档文本内容（将提交给AI）：")
            print("=" * 100)
            print(text_preview)
            print("=" * 100)
            logger.info("=" * 100)
            logger.info("[读取Word文档] 📄 提取的文档文本内容（将提交给AI）：")
            logger.info("=" * 100)
            logger.info(text_preview)
            logger.info("=" * 100)
            
            # 直接将文档文本传给AI，让AI分析并生成格式化指令
            print("=" * 100)
            print(f"[读取Word文档] 步骤2/2: 将文档传给AI分析并生成格式化指令...")
            print(f"  开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            print(f"  文档文本长度: {text_length} 字符")
            print(f"  即将调用AI模型，这可能需要一些时间（通常30-120秒）...")
            print("=" * 100)
            import sys
            sys.stdout.flush()
            logger.info("=" * 100)
            logger.info(f"[读取Word文档] 步骤2/2: 将文档传给AI分析并生成格式化指令...")
            logger.info(f"  开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            logger.info(f"  文档文本长度: {text_length} 字符")
            logger.info("=" * 100)
            
            try:
                print(f"[读取Word文档] 正在调用AI分析服务...")
                logger.info(f"[读取Word文档] 正在调用AI分析服务...")
                ai_start_time = time.time()
                format_result = await cls._analyze_format_with_ai(
                    query_db,
                    document_text,
                    config_id
                )
                ai_elapsed = time.time() - ai_start_time
                format_instructions = format_result['json_instructions']
                natural_language_description = format_result['natural_language_description']
                print("=" * 100)
                print(f"[读取Word文档] ✓ 步骤2/2: AI分析完成")
                print(f"  完成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                print(f"  耗时: {ai_elapsed:.2f} 秒 ({ai_elapsed / 60:.2f} 分钟)")
                print(f"  自然语言描述长度: {len(natural_language_description)} 字符")
                print(f"  JSON格式指令长度: {len(format_instructions)} 字符")
                print("=" * 100)
                sys.stdout.flush()
                logger.info("=" * 100)
                logger.info(f"[读取Word文档] ✓ 步骤2/2: AI分析完成")
                logger.info(f"  完成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                logger.info(f"  耗时: {ai_elapsed:.2f} 秒 ({ai_elapsed / 60:.2f} 分钟)")
                logger.info(f"  自然语言描述长度: {len(natural_language_description)} 字符")
                logger.info(f"  JSON格式指令长度: {len(format_instructions)} 字符")
                logger.info("=" * 100)
                
                # 校验格式指令（格式规范、一致性、数据质量）
                # 注意：校验在AI分析完成后进行，不单独作为步骤
                print(f"[读取Word文档] 校验格式指令...")
                logger.info(f"[读取Word文档] 校验格式指令...")
                
                try:
                    # 解析JSON指令
                    format_instruction_data = json.loads(format_instructions) if isinstance(format_instructions, str) else format_instructions
                    
                    # 读取完整指令系统用于校验
                    universal_instruction_system = await cls._get_universal_instruction_system(query_db)
                    
                    if universal_instruction_system:
                        # 执行校验
                        validation_result = cls._validate_instruction_system(
                            natural_language_description,
                            format_instruction_data,
                            universal_instruction_system
                        )
                        
                        if not validation_result['valid']:
                            error_msg = f"格式指令校验失败: {', '.join(validation_result['errors'][:5])}"  # 只显示前5个错误
                            logger.error(error_msg)
                            print(f"[读取Word文档] ⚠ 校验警告: {error_msg}")
                            # 不抛出异常，只记录警告，允许继续保存（因为数据质量校验已经修正了异常值）
                        else:
                            print(f"[读取Word文档] ✓ 格式指令校验通过")
                            logger.info(f"[读取Word文档] ✓ 格式指令校验通过")
                    else:
                        logger.warning("未找到完整指令系统，跳过格式规范校验和一致性校验")
                        print(f"[读取Word文档] ⚠ 未找到完整指令系统，跳过格式规范校验和一致性校验")
                        
                except Exception as e:
                    logger.warning(f"格式指令校验过程出错: {str(e)}，继续保存指令")
                    print(f"[读取Word文档] ⚠ 格式指令校验过程出错，继续保存指令")
                
            except Exception as e:
                print("=" * 100)
                print(f"[读取Word文档] ✗ 步骤2/2: AI分析失败")
                print(f"  错误: {str(e)}")
                import traceback
                print(traceback.format_exc())
                print("=" * 100)
                sys.stdout.flush()
                logger.error("=" * 100)
                logger.error(f"[读取Word文档] ✗ 步骤2/2: AI分析失败")
                logger.error(f"  错误: {str(e)}")
                logger.error("=" * 100, exc_info=True)
                raise
            
            inst_len = len(format_instructions)
            inst_preview = format_instructions[:200] if format_instructions else 'N/A'
            desc_len = len(natural_language_description)
            desc_preview = natural_language_description[:200] if natural_language_description else 'N/A'
            print(f"[读取Word文档] 完成 - 文件: {word_file_path}")
            print(f"  自然语言描述长度: {desc_len} 字符")
            print(f"  自然语言描述前200字符: {desc_preview}")
            print(f"  JSON格式指令长度: {inst_len} 字符")
            print(f"  JSON格式指令前200字符: {inst_preview}")
            logger.info(f"[读取Word文档] 完成 - 文件: {word_file_path}")
            logger.info(f"  自然语言描述长度: {desc_len} 字符")
            logger.info(f"  自然语言描述前200字符: {desc_preview}")
            logger.info(f"  JSON格式指令长度: {inst_len} 字符")
            logger.info(f"  JSON格式指令前200字符: {inst_preview}")
            
            return {
                'format_instructions': format_instructions,  # JSON格式指令（用于执行）
                'natural_language_description': natural_language_description,  # 自然语言描述（用于展示）
                'document_text': document_text,  # 文档文本内容
                'file_path': word_file_path
            }
            
        except ServiceException as e:
            # ServiceException 直接重新抛出，保留原始错误信息
            error_msg = e.message if hasattr(e, 'message') else str(e)
            print("=" * 100)
            print(f"[读取Word文档] ✗ ServiceException: {error_msg}")
            print("=" * 100)
            import sys
            sys.stdout.flush()
            logger.error(f"[读取Word文档] ServiceException: {error_msg}", exc_info=True)
            raise
        except Exception as e:
            # 获取完整的错误信息
            error_msg = str(e) if str(e) else repr(e)
            error_type = type(e).__name__
            print("=" * 100)
            print(f"[读取Word文档] ✗ 异常: {error_type}")
            print(f"  错误信息: {error_msg}")
            import traceback
            print(traceback.format_exc())
            print("=" * 100)
            sys.stdout.flush()
            logger.error(f"[读取Word文档] 失败: {error_type} - {error_msg}", exc_info=True)
            raise ServiceException(message=f'读取Word文档失败: {error_msg}')
    
    @staticmethod
    def _extract_document_text(doc: Any) -> str:
        """
        提取Word文档的文本内容（保留所有文字信息，不丢失任何文字）
        
        重要原则：
        1. 所有有文字的段落都必须保留
        2. 空行可以忽略，但空行前后的文字不能丢失
        3. 确保"空两行"等文字说明完整保留
        
        :param doc: python-docx Document对象
        :return: 文档文本内容（保留所有文字信息）
        """
        paragraphs_text = []
        for para in doc.paragraphs:
            text = para.text  # 获取原始文本（不strip，保留前后空格）
            # 只要有文字（即使只有空格），都要保留
            # 完全空的段落可以忽略（但保留换行以维持结构）
            if text or text.strip():  # 有文字或只有空格都保留
                paragraphs_text.append(text)
            else:
                # 完全空的段落，保留一个空行以维持文档结构
                paragraphs_text.append('')
        
        # 合并段落，用换行符分隔
        result = '\n'.join(paragraphs_text)
        
        # 验证：确保没有丢失文字
        total_chars = sum(len(para.text) for para in doc.paragraphs)
        extracted_chars = len(result.replace('\n', ''))
        if extracted_chars < total_chars * 0.9:  # 如果提取的字符少于90%，可能有问题
            logger.warning(f"文本提取可能不完整：原始字符数={total_chars}，提取字符数={extracted_chars}")
        
        return result
    
    @classmethod
    def _extract_document_content(cls, doc: Any) -> Dict[str, Any]:
        """
        提取Word文档的内容和格式信息
        
        识别以下格式要求：
        1. 字体格式：字体名称、大小、颜色、粗体、斜体、下划线
        2. 段落格式：对齐方式、行距、段前距、段后距、首行缩进、左右缩进
        3. 标题格式：各级标题的样式（字体、大小、对齐等）
        4. 页面设置：页边距、纸张大小、页眉页脚
        5. 样式信息：文档中使用的样式名称和属性
        
        :param doc: python-docx Document对象
        :return: 文档内容字典
        """
        content = {
            'paragraphs': [],
            'styles': {},
            'headings': {},  # 识别标题格式
            'format_info': {
                'font_name': None,
                'font_size': None,
                'line_spacing': None,
                'paragraph_spacing': None,
                'margins': None,
                'page_size': None,
            }
        }
        
        # 提取段落内容和格式
        logger.debug(f"[提取文档内容] 开始提取段落，总段落数: {len(doc.paragraphs)}")
        for idx, para in enumerate(doc.paragraphs):
            if idx < 5 or idx % 50 == 0:  # 前5个段落和每50个段落记录一次
                logger.debug(f"[提取文档内容] 处理段落 {idx+1}/{len(doc.paragraphs)}: {para.text[:50] if para.text else '(空段落)'}")
            para_info = {
                'text': para.text,
                'style': para.style.name if para.style else None,
                'alignment': str(para.alignment) if para.alignment else None,
            }
            
            # 提取段落格式信息
            if para.paragraph_format:
                para_info['paragraph_format'] = {
                    'alignment': str(para.alignment) if para.alignment else None,
                    'line_spacing': str(para.paragraph_format.line_spacing) if para.paragraph_format.line_spacing else None,
                    'line_spacing_rule': str(para.paragraph_format.line_spacing_rule) if para.paragraph_format.line_spacing_rule else None,
                    'space_before': str(para.paragraph_format.space_before) if para.paragraph_format.space_before else None,
                    'space_after': str(para.paragraph_format.space_after) if para.paragraph_format.space_after else None,
                    'first_line_indent': str(para.paragraph_format.first_line_indent) if para.paragraph_format.first_line_indent else None,
                    'left_indent': str(para.paragraph_format.left_indent) if para.paragraph_format.left_indent else None,
                    'right_indent': str(para.paragraph_format.right_indent) if para.paragraph_format.right_indent else None,
                }
            
            # 提取段落中的文本运行格式信息
            runs_info = []
            for run in para.runs:
                run_info = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name if run.font.name else None,
                    'font_size': float(run.font.size.pt) if run.font.size and run.font.size.pt else None,  # 转换为磅值（浮点数）
                    'font_color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                }
                runs_info.append(run_info)
            para_info['runs'] = runs_info
            content['paragraphs'].append(para_info)
            
            # 识别标题格式（通过样式名称或格式特征）
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if 'heading' in style_name or 'title' in style_name or '标题' in style_name:
                    # 提取标题级别和格式
                    heading_level = None
                    if 'heading 1' in style_name or '标题 1' in style_name or 'h1' in style_name:
                        heading_level = 1
                    elif 'heading 2' in style_name or '标题 2' in style_name or 'h2' in style_name:
                        heading_level = 2
                    elif 'heading 3' in style_name or '标题 3' in style_name or 'h3' in style_name:
                        heading_level = 3
                    
                    if heading_level:
                        if heading_level not in content['headings']:
                            # 提取完整的标题格式信息
                            heading_info = {
                                'style_name': para.style.name,
                                'font_name': runs_info[0]['font_name'] if runs_info else None,
                                'font_size': float(runs_info[0]['font_size']) if runs_info and runs_info[0].get('font_size') is not None else None,
                                'bold': runs_info[0]['bold'] if runs_info else None,
                                'alignment': para_info['alignment'],
                            }
                            # 添加段落格式信息
                            if para.paragraph_format:
                                if para.paragraph_format.space_before:
                                    heading_info['spacing_before'] = str(para.paragraph_format.space_before)
                                if para.paragraph_format.space_after:
                                    heading_info['spacing_after'] = str(para.paragraph_format.space_after)
                            
                            # 计算标题前后的空行数
                            spacing_before_lines = cls._count_empty_lines_before(doc.paragraphs, idx)
                            spacing_after_lines = cls._count_empty_lines_after(doc.paragraphs, idx)
                            if spacing_before_lines > 0:
                                heading_info['spacing_before_lines'] = spacing_before_lines
                            if spacing_after_lines > 0:
                                heading_info['spacing_after_lines'] = spacing_after_lines
                            
                            content['headings'][f'h{heading_level}'] = heading_info
            
            # 识别特殊格式（目录、摘要、关键词、结论等）
            para_text_lower = para.text.strip().lower()
            if '目录' in para.text or 'table of contents' in para_text_lower:
                if 'table_of_contents' not in content:
                    content['table_of_contents'] = {
                        'title_para': para_info.copy(),
                        'entries': []
                    }
                    # 计算目录标题前后的空行数
                    title_before_lines = cls._count_empty_lines_before(doc.paragraphs, idx)
                    title_after_lines = cls._count_empty_lines_after(doc.paragraphs, idx)
                    if title_before_lines > 0:
                        content['table_of_contents']['title_before_lines'] = title_before_lines
                    if title_after_lines > 0:
                        content['table_of_contents']['title_after_lines'] = title_after_lines
                # 标记进入目录区域
                content['_in_toc_section'] = True
            elif 'abstract' in para_text_lower or '摘要' in para.text:
                # 目录区域结束（遇到摘要）
                if '_in_toc_section' in content:
                    content['_in_toc_section'] = False
                    
            # 识别目录条目（在目录区域内，且包含页码或点线）
            if content.get('_in_toc_section', False):
                # 目录条目的特征：
                # 1. 包含页码（数字结尾，可能有前导点）
                # 2. 或包含点线（多个点或下划线）
                # 3. 或包含章节编号（一、二、三、或1、2、3等）
                # 4. 排除"目录"标题本身和"摘要"等章节标题
                import re
                para_text = para.text.strip()
                
                # 排除章节标题（摘要、关键词、结论等）
                exclude_keywords = ['摘要', '关键词', 'Abstract', 'Key words', '结论', '参考文献', '附录', '致谢']
                if any(keyword in para_text for keyword in exclude_keywords):
                    # 遇到这些关键词，说明目录区域结束
                    content['_in_toc_section'] = False
                else:
                    toc_patterns = [
                        r'^[一二三四五六七八九十]+[、．]',  # 中文编号（一、二、三、）
                        r'^[（(][一二三四五六七八九十]+[）)]',  # 中文编号（（一）、（二））
                        r'^\d+[、．]',  # 数字编号（1、2、3、）
                        r'^\d+[.、]',  # 数字编号（1. 2. 3.）
                        r'[．…]+',  # 点线（...）
                        r'\d+$',  # 以数字结尾（页码）
                        r'^[A-Z]\.',  # 字母编号（A. B. C.）
                    ]
                    is_toc_entry = False
                    for pattern in toc_patterns:
                        if re.search(pattern, para_text):
                            is_toc_entry = True
                            break
                    
                    if is_toc_entry and 'table_of_contents' in content:
                        # 提取目录条目的格式信息
                        # 计算缩进值（用于区分级别）
                        left_indent = 0
                        if para.paragraph_format and para.paragraph_format.left_indent:
                            try:
                                left_indent = int(para.paragraph_format.left_indent)
                            except:
                                left_indent = 0
                        
                        # 判断目录级别（根据缩进值）
                        toc_level = 1
                        if left_indent > 0:
                            if left_indent <= 200:  # 约2字符
                                toc_level = 2
                            elif left_indent <= 400:  # 约4字符
                                toc_level = 3
                        
                        toc_entry_info = {
                            'text': para_text,
                            'level': toc_level,
                            'para_format': para_info.get('paragraph_format', {}),
                            'runs': runs_info,
                            # 提取目录条目的字体和大小（从runs中提取，取第一个run的格式）
                            'font_name': runs_info[0]['font_name'] if runs_info else None,
                            'font_size': float(runs_info[0]['font_size']) if runs_info and runs_info[0].get('font_size') is not None else None,
                            'alignment': para_info.get('alignment'),
                            'left_indent': left_indent,
                            'line_spacing': para_info.get('paragraph_format', {}).get('line_spacing'),
                        }
                        content['table_of_contents']['entries'].append(toc_entry_info)
            elif '摘要' in para.text or 'abstract' in para_text_lower:
                if 'abstract' not in content:
                    content['abstract'] = {
                        'title_para': para_info.copy(),
                        'content_paras': []
                    }
            elif '关键词' in para.text or 'keywords' in para_text_lower or 'key words' in para_text_lower:
                if 'keywords' not in content:
                    content['keywords'] = {
                        'para': para_info.copy()
                    }
            elif '结论' in para.text or 'conclusion' in para_text_lower:
                if 'conclusion' not in content:
                    content['conclusion'] = {
                        'title_para': para_info.copy(),
                        'content_paras': []
                    }
                    # 计算结论标题前后的空行数
                    title_before_lines = cls._count_empty_lines_before(doc.paragraphs, idx)
                    title_after_lines = cls._count_empty_lines_after(doc.paragraphs, idx)
                    if title_before_lines > 0:
                        content['conclusion']['title_before_lines'] = title_before_lines
                    if title_after_lines > 0:
                        content['conclusion']['title_after_lines'] = title_after_lines
            elif '参考文献' in para.text or 'references' in para_text_lower or '参 考 文 献' in para.text:
                if 'references' not in content:
                    content['references'] = {
                        'title_para': para_info.copy(),
                        'entries': []
                    }
                    # 计算参考文献标题前后的空行数
                    title_before_lines = cls._count_empty_lines_before(doc.paragraphs, idx)
                    title_after_lines = cls._count_empty_lines_after(doc.paragraphs, idx)
                    if title_before_lines > 0:
                        content['references']['title_before_lines'] = title_before_lines
                    if title_after_lines > 0:
                        content['references']['title_after_lines'] = title_after_lines
            elif '致谢' in para.text or '致　谢' in para.text or 'acknowledgement' in para_text_lower:
                if 'acknowledgement' not in content:
                    content['acknowledgement'] = {
                        'title_para': para_info.copy(),
                        'content_paras': []
                    }
                    # 计算致谢标题前后的空行数
                    title_before_lines = cls._count_empty_lines_before(doc.paragraphs, idx)
                    title_after_lines = cls._count_empty_lines_after(doc.paragraphs, idx)
                    if title_before_lines > 0:
                        content['acknowledgement']['title_before_lines'] = title_before_lines
                    if title_after_lines > 0:
                        content['acknowledgement']['title_after_lines'] = title_after_lines
        
        # 提取默认样式信息（从正文段落，取多个段落进行统计分析）
        body_paragraphs = [p for p in doc.paragraphs if p.style and 'heading' not in p.style.name.lower() and 'title' not in p.style.name.lower()]
        if body_paragraphs:
            # 统计最常见的字体和格式（取前3个正文段落）
            sample_paragraphs = body_paragraphs[:3]
            font_names = []
            font_sizes = []
            
            for para in sample_paragraphs:
                if para.runs:
                    for run in para.runs:
                        if run.font.name:
                            font_names.append(run.font.name)
                        if run.font.size and run.font.size.pt:
                            font_sizes.append(float(run.font.size.pt))  # 转换为磅值（浮点数）
            
            # 使用最常见的字体和字号
            if font_names:
                from collections import Counter
                most_common_font = Counter(font_names).most_common(1)[0][0]
                content['format_info']['font_name'] = most_common_font
            if font_sizes:
                from collections import Counter
                most_common_size = Counter(font_sizes).most_common(1)[0][0]
                content['format_info']['font_size'] = most_common_size
            
            # 提取段落格式（从第一个正文段落，但记录所有段落的格式信息）
            para_formats = []
            for para in sample_paragraphs:
                if para.paragraph_format:
                    para_format = {}
                    if para.paragraph_format.line_spacing:
                        para_format['line_spacing'] = str(para.paragraph_format.line_spacing)
                    if para.paragraph_format.space_before:
                        para_format['space_before'] = str(para.paragraph_format.space_before)
                    if para.paragraph_format.space_after:
                        para_format['space_after'] = str(para.paragraph_format.space_after)
                    if para.paragraph_format.first_line_indent:
                        para_format['first_line_indent'] = str(para.paragraph_format.first_line_indent)
                    if para.paragraph_format.left_indent:
                        para_format['left_indent'] = str(para.paragraph_format.left_indent)
                    if para.paragraph_format.right_indent:
                        para_format['right_indent'] = str(para.paragraph_format.right_indent)
                    if para_format:
                        para_formats.append(para_format)
            
            # 使用第一个段落的格式作为默认格式
            if para_formats:
                content['format_info'].update(para_formats[0])
        
        # 提取节信息（页边距、纸张大小等）
        logger.debug(f"[提取文档内容] 开始提取页面设置信息...")
        if doc.sections:
            section = doc.sections[0]
            content['format_info']['margins'] = {
                'top': str(section.top_margin),
                'bottom': str(section.bottom_margin),
                'left': str(section.left_margin),
                'right': str(section.right_margin),
            }
            logger.debug(f"[提取文档内容] 页边距: 上={section.top_margin}, 下={section.bottom_margin}, 左={section.left_margin}, 右={section.right_margin}")
            
            # 提取纸张大小
            if section.page_width and section.page_height:
                content['format_info']['page_size'] = {
                    'width': str(section.page_width),
                    'height': str(section.page_height),
                }
                logger.debug(f"[提取文档内容] 页面大小: 宽={section.page_width}, 高={section.page_height}")
        
        # 提取文档样式信息
        logger.debug(f"[提取文档内容] 开始提取样式信息...")
        if hasattr(doc.styles, 'styles'):
            for style in doc.styles.styles:
                if style.name and style.name not in ['Normal', 'Default Paragraph Font']:
                    content['styles'][style.name] = {
                        'type': str(style.type) if hasattr(style, 'type') else None,
                    }
        logger.debug(f"[提取文档内容] 样式提取完成，识别到 {len(content['styles'])} 个样式")
        
        logger.info(f"[提取文档内容] 文档内容提取完成")
        logger.info(f"  段落数: {len(content['paragraphs'])}")
        logger.info(f"  标题数: {len(content['headings'])}")
        logger.info(f"  样式数: {len(content['styles'])}")
        if content.get('table_of_contents'):
            logger.info(f"  目录条目数: {len(content.get('table_of_contents', {}).get('entries', []))}")
        
        return content
    
    @staticmethod
    def _count_empty_lines_before(paragraphs: list, current_idx: int) -> int:
        """
        计算指定段落之前的连续空行数
        
        :param paragraphs: 段落列表
        :param current_idx: 当前段落索引
        :return: 空行数
        """
        if current_idx <= 0:
            return 0
        
        empty_count = 0
        # 从当前段落向前查找连续的空段落
        for i in range(current_idx - 1, -1, -1):
            para = paragraphs[i]
            # 判断是否为空段落：文本为空或只包含空白字符
            if not para.text or para.text.strip() == '':
                empty_count += 1
            else:
                # 遇到非空段落，停止计数
                break
        
        return empty_count
    
    @staticmethod
    def _count_empty_lines_after(paragraphs: list, current_idx: int) -> int:
        """
        计算指定段落之后的连续空行数
        
        :param paragraphs: 段落列表
        :param current_idx: 当前段落索引
        :return: 空行数
        """
        if current_idx >= len(paragraphs) - 1:
            return 0
        
        empty_count = 0
        # 从当前段落向后查找连续的空段落
        for i in range(current_idx + 1, len(paragraphs)):
            para = paragraphs[i]
            # 判断是否为空段落：文本为空或只包含空白字符
            if not para.text or para.text.strip() == '':
                empty_count += 1
            else:
                # 遇到非空段落，停止计数
                break
        
        return empty_count
    
    @classmethod
    async def _analyze_format_with_ai(
        cls,
        query_db: AsyncSession,
        document_text: str,
        config_id: Optional[int] = None
    ) -> Dict[str, str]:
        """
        使用AI直接分析文档文本并生成格式化指令
        
        优化后的流程：直接将文档文本传给AI，让AI一步完成：
        1. 理解格式要求
        2. 生成自然语言描述
        3. 生成子集指令系统
        
        :param query_db: 数据库会话
        :param document_text: 文档文本内容
        :param config_id: AI模型配置ID（可选）
        :return: 包含自然语言描述和JSON指令的字典
            {
                'natural_language_description': '自然语言格式描述（用于展示）',
                'json_instructions': 'JSON格式指令（用于执行）'
            }
        """
        print("[AI格式分析] 开始分析文档格式...")
        logger.info("[AI格式分析] 开始分析文档格式...")
        
        # 1. 读取完整指令系统
        print("[AI格式分析] 读取完整指令系统...")
        logger.info("[AI格式分析] 读取完整指令系统...")
        universal_instruction_system = await cls._get_universal_instruction_system(query_db)
        
        if not universal_instruction_system:
            logger.warning("未找到完整指令系统，将使用简化方法生成指令")
            # 如果没有完整指令系统，使用简化方法
            return await cls._analyze_format_with_ai_simple(query_db, document_text, config_id)
        
        print(f"[AI格式分析] ✓ 完整指令系统读取完成（版本: {universal_instruction_system.get('version', 'N/A')}）")
        logger.info(f"[AI格式分析] ✓ 完整指令系统读取完成（版本: {universal_instruction_system.get('version', 'N/A')}）")
        
        # 2. 直接将文档文本传给AI，让AI分析并生成格式化指令
        print("[AI格式分析] 将文档传给AI分析并生成格式化指令...")
        logger.info("[AI格式分析] 将文档传给AI分析并生成格式化指令...")
        result = await cls._generate_format_instructions_directly(
            query_db,
            document_text,
            universal_instruction_system,
            config_id
        )
        natural_language = result['natural_language_description']
        subset_instruction = result['subset_instruction']
        print(f"[AI格式分析] ✓ AI分析完成")
        logger.info(f"[AI格式分析] ✓ AI分析完成")
        
        # 3. 打印生成的内容
        print("=" * 100)
        print("[AI格式分析] 📄 AI生成的格式化内容：")
        print("=" * 100)
        print("【自然语言格式描述】")
        print(natural_language[:500] + "..." if len(natural_language) > 500 else natural_language)
        print("=" * 100)
        print("【JSON格式指令】")
        subset_instruction_json = json.dumps(subset_instruction, ensure_ascii=False, indent=2)
        print(subset_instruction_json[:500] + "..." if len(subset_instruction_json) > 500 else subset_instruction_json)
        print("=" * 100)
        import sys
        sys.stdout.flush()
        
        logger.info("=" * 100)
        logger.info("[AI格式分析] 📄 AI生成的格式化内容：")
        logger.info("=" * 100)
        logger.info("【自然语言格式描述】")
        logger.info(natural_language[:500] + "..." if len(natural_language) > 500 else natural_language)
        logger.info("=" * 100)
        logger.info("【JSON格式指令】")
        logger.info(subset_instruction_json[:500] + "..." if len(subset_instruction_json) > 500 else subset_instruction_json)
        logger.info("=" * 100)
        
        # 3. 验证和修正格式指令（动态修正异常值）
        print("[AI格式分析] 验证和修正格式指令...")
        logger.info("[AI格式分析] 验证和修正格式指令...")
        
        try:
            # 立即验证和修正（动态修正异常值）
            validated_config = cls._validate_and_fix_format_config(subset_instruction)
            subset_instruction = validated_config
            
            print(f"[AI格式分析] ✓ 格式指令验证和修正完成")
            logger.info(f"[AI格式分析] ✓ 格式指令验证和修正完成")
        except Exception as e:
            logger.warning(f"格式指令验证失败: {str(e)}，使用原始指令")
            print(f"[AI格式分析] ⚠ 格式指令验证失败，使用原始指令")
        
        print(f"[AI格式分析] ✓ 格式分析完成")
        logger.info(f"[AI格式分析] ✓ 格式分析完成")
        
        return {
            'natural_language_description': natural_language,
            'json_instructions': json.dumps(subset_instruction, ensure_ascii=False, indent=2)
        }
    
    @classmethod
    async def _analyze_format_with_ai_legacy(
        cls,
        query_db: AsyncSession,
        document_content: Dict[str, Any],
        config_id: Optional[int] = None
    ) -> Dict[str, str]:
        """
        原有的AI分析方法（作为回退方案）
        
        :param query_db: 数据库会话
        :param document_content: 文档内容
        :param config_id: AI模型配置ID（可选）
        :return: 包含自然语言描述和JSON指令的字典
        """
        # 构建提示词
        print("[AI格式分析] 使用原有方法生成指令...")
        logger.info("[AI格式分析] 使用原有方法生成指令...")
        prompt = cls._build_format_analysis_prompt(document_content)
        prompt_len = len(prompt)
        
        # 获取AI提供商
        llm_provider, model_config = await AiGenerationService._get_ai_provider(query_db, config_id)
        
        # 调用AI分析
        messages = [
            {
                "role": "system",
                "content": "你是一位专业的文档格式分析专家，擅长分析Word文档的格式要求并生成详细的格式化指令。\n\n重要：你需要同时生成两个输出：\n1. 自然语言格式描述（用于展示给用户）\n2. JSON格式指令（用于系统执行）\n\n格式化指令是指令性的、可执行的格式规范，它描述了如何将格式应用到文档中。指令中包含格式信息，但更重要的是，它是指令性的。"
            },
            {"role": "user", "content": prompt}
        ]
        
        try:
            import asyncio
            response = await asyncio.wait_for(
                llm_provider.chat(messages, temperature=0.3, max_tokens=4000),
                timeout=120.0
            )
            
            # 解析AI响应
            result = cls._parse_ai_format_response(response)
            
            # 验证和修正
            try:
                format_config = json.loads(result['json_instructions']) if isinstance(result['json_instructions'], str) else result['json_instructions']
                validated_config = cls._validate_and_fix_format_config(format_config)
                result['json_instructions'] = json.dumps(validated_config, ensure_ascii=False, indent=2)
            except Exception as e:
                logger.warning(f"格式指令验证失败: {str(e)}，使用原始指令")
            
            return result
            
        except Exception as e:
            logger.error(f"AI分析失败: {str(e)}", exc_info=True)
            raise ServiceException(message=f"AI分析失败: {str(e)}")
    
    @classmethod
    def _parse_ai_format_response(cls, response: str) -> Dict[str, str]:
        """
        解析AI响应，提取自然语言描述和JSON指令
        
        支持多种AI响应格式：
        1. 【格式要求描述】...【格式化指令】...```json {...} ```
        2. 自然语言描述...```json {...} ```
        3. 纯JSON格式
        4. 纯自然语言格式
        
        :param response: AI返回的完整响应
        :return: 包含自然语言描述和JSON指令的字典
        """
        import re
        
        response = response.strip()
        
        # 尝试提取JSON部分（可能在markdown代码块中，也可能直接是JSON）
        json_instructions = None
        natural_language_description = ""
        
        # 方法1：查找JSON代码块（支持嵌套JSON）
        # 匹配 ```json ... ``` 或 ``` ... ``` 格式
        json_code_block_pattern = r'```(?:json)?\s*(\{.*?\})\s*```'
        json_match = re.search(json_code_block_pattern, response, re.DOTALL)
        if json_match:
            json_candidate = json_match.group(1).strip()
            # 验证JSON是否有效
            try:
                json.loads(json_candidate)
                json_instructions = json_candidate
                # 自然语言描述是JSON代码块之前的内容
                natural_language_description = response[:json_match.start()].strip()
                # 移除可能的标记（如【格式要求描述】、【格式化指令】等）
                natural_language_description = re.sub(r'【.*?】', '', natural_language_description).strip()
            except json.JSONDecodeError:
                json_instructions = None
        
        # 方法2：如果方法1失败，查找大括号包围的JSON对象（支持嵌套）
        if not json_instructions:
            # 从后往前查找，找到最后一个完整的JSON对象
            brace_count = 0
            json_start = -1
            json_end = -1
            
            for i in range(len(response) - 1, -1, -1):
                if response[i] == '}':
                    if brace_count == 0:
                        json_end = i + 1
                    brace_count += 1
                elif response[i] == '{':
                    brace_count -= 1
                    if brace_count == 0:
                        json_start = i
                        json_str = response[json_start:json_end]
                        try:
                            # 验证是否是有效的JSON
                            json.loads(json_str)
                            json_instructions = json_str
                            # 自然语言描述是JSON之前的内容
                            natural_language_description = response[:json_start].strip()
                            # 移除可能的标记
                            natural_language_description = re.sub(r'【.*?】', '', natural_language_description).strip()
                            break
                        except json.JSONDecodeError:
                            continue
        
        # 方法3：如果找不到JSON，尝试直接解析整个响应为JSON
        if not json_instructions:
            try:
                json.loads(response)
                # 如果整个响应是JSON，则没有自然语言描述
                json_instructions = response
                natural_language_description = "格式要求已提取为JSON指令格式。"
            except json.JSONDecodeError:
                # 如果无法解析，将整个响应作为自然语言描述，JSON为空
                natural_language_description = response
                # 移除可能的标记
                natural_language_description = re.sub(r'【.*?】', '', natural_language_description).strip()
                json_instructions = ""
                logger.warning("无法从AI响应中提取JSON指令，将整个响应作为自然语言描述")
        
        # 如果自然语言描述为空或太短，生成一个默认描述
        if not natural_language_description or len(natural_language_description.strip()) < 10:
            natural_language_description = "已从格式文件中提取格式要求，并生成格式化指令。"
        
        # 验证JSON是否有效，如果无效则尝试修复
        if json_instructions:
            try:
                json.loads(json_instructions)
            except json.JSONDecodeError:
                logger.warning("提取的JSON指令无效，尝试修复...")
                # 尝试修复常见的JSON问题
                try:
                    extracted_json = cls._extract_json_from_text(json_instructions)
                    if isinstance(extracted_json, dict):
                        json_instructions = json.dumps(extracted_json, ensure_ascii=False)
                    elif isinstance(extracted_json, str):
                        json_instructions = extracted_json
                    else:
                        json_instructions = str(extracted_json)
                except Exception as e:
                    logger.error(f"修复JSON失败: {str(e)}")
                    json_instructions = ""
        
        return {
            'natural_language_description': natural_language_description,
            'json_instructions': json_instructions if json_instructions else ""
        }
    
    @classmethod
    def _build_format_analysis_prompt(cls, document_content: Dict[str, Any]) -> str:
        """构建格式分析提示词（方案B：要求同时生成自然语言描述和JSON指令）"""
        # 将文档内容转换为JSON字符串（用于AI分析）
        # 限制长度，避免超出token限制，但保留关键格式信息
        content_json = json.dumps(document_content, ensure_ascii=False, indent=2)
        if len(content_json) > 3000:
            # 如果内容太长，只保留格式信息部分
            content_dict = document_content.copy()
            # 保留段落格式信息，但截断文本内容
            if 'paragraphs' in content_dict:
                for para in content_dict['paragraphs']:
                    if 'text' in para and len(para['text']) > 100:
                        para['text'] = para['text'][:100] + '...'
            content_json = json.dumps(content_dict, ensure_ascii=False, indent=2)[:3000]
        
        prompt = f"""你是文档格式分析专家。分析Word文档格式并生成格式化指令。

## ⚠️ 核心原则：
1. 严格按文档实际格式提取，不使用默认值或猜测
2. 原样返回：字体名称、字号、行距、缩进等必须与文档一致

## 文档格式信息：
{content_json}

## 任务：
生成两个输出：
1. **自然语言格式描述**：用自然语言描述格式要求（字体、字号、行距、页边距、标题格式等）
2. **JSON格式指令**：严格按照JSON结构返回格式化指令

## 格式提取要求：

### 1. 字体格式
- 默认字体、字号、颜色（从正文段落提取）

### 2. 段落格式
- 对齐方式（left/center/right/justify）
- 行距倍数（如1.5）
- 段前距、段后距（磅）
- 首行缩进（磅，如24表示2字符）
- 左右缩进（磅）

### 3. 标题格式（h1/h2/h3）
- 字体名称、字号、是否加粗、对齐方式、段前距、段后距

### 4. 页面设置
- 页边距（磅，1英寸=72磅，1厘米=28.35磅）
- 纸张大小（A4/Letter等）

### 5. 目录格式（如存在）
- 目录标题格式（字体、字号、对齐）
- 目录条目格式（多级，字体、字号、行距、对齐、缩进）

### 6. 特殊格式（如存在）
- **目录格式**：如果文档中有目录，必须详细识别目录的格式要求
  - **目录标题格式**：
    - 字体名称、字体大小、是否加粗、对齐方式
    - 例如："目录"可能是小3号黑体，居中
  - **目录条目格式**（非常重要）：
    - 必须从文档中实际提取目录条目的格式
    - 字体名称：目录条目的字体（如：宋体、仿宋等）
    - 字体大小：目录条目的字号（如：12磅表示小四号，14磅表示四号）
    - 行距：目录条目的行距倍数（如：1.5倍）
    - 对齐方式：目录条目的对齐方式（如：justify表示分散对齐，left表示左对齐）
    - 段前距、段后距：目录条目的段间距
    - **注意**：目录条目可能包含多级标题（一级、二级、三级），需要识别不同级别的格式
    - **注意**：目录条目通常包含点线（...）或页码，需要识别这些特殊格式

## 重要格式要求（必须从文档中实际提取）：

### 1. 标题格式识别（必须从文档中实际提取）
- **一级标题（h1）**：
  - 字体：从文档中实际提取（可能是黑体、楷体、宋体加粗等，不同学校可能有不同要求）
  - 大小：从文档中实际提取（准确识别实际磅值）
  - 加粗：从文档中实际提取（true/false，必须准确识别）
  - 对齐：从文档中实际提取（left/center/right）
  
- **二级标题（h2）**：
  - 字体：从文档中实际提取
  - 大小：从文档中实际提取
  - 加粗：从文档中实际提取
  - 对齐：从文档中实际提取
  
- **三级标题（h3）**：
  - 字体：从文档中实际提取
  - 大小：从文档中实际提取
  - 加粗：从文档中实际提取
  - 对齐：从文档中实际提取

**重要**：不同学校可能有不同的格式要求，必须从文档中实际提取，不要使用标准值或猜测。
例如：有些学校标题用黑体，有些用楷体；有些标题加粗，有些不加粗；字号也可能不同。

### 2. 正文格式识别（必须从文档中实际提取）
- **默认字体**：
  - 字体名称：从文档中实际提取（可能是宋体、楷体、Times New Roman等）
  - 字体大小：从文档中实际提取（准确识别实际磅值，如12、14、16等）
  - 字体颜色：从文档中实际提取（通常是黑色000000）
  
- **段落格式**：
  - 行距：从文档中实际提取（准确识别实际倍数，如1.0、1.5、2.0等）
  - 首行缩进：从文档中实际提取（准确识别实际磅值，可能是0、24、28等）
  - 对齐方式：从文档中实际提取（left/center/right/justify）
  - 右缩进：从文档中实际提取（通常是0，但必须准确识别）

**重要**：不同学校可能有不同的格式要求，必须从文档中实际提取，不要使用标准值或猜测。
例如：有些学校行距是1.5倍，有些是1.0倍或2.0倍；有些学校需要首行缩进，有些不需要。
- **摘要格式**：如果文档中有摘要，识别摘要的格式
  - 摘要标题格式
  - 摘要正文格式（通常是小4号宋体）
- **关键词格式**：识别关键词的格式要求
- **结论格式**：识别结论部分的格式要求
- **表格格式**：如果文档中有表格，识别表格的边框、对齐等格式
- **列表格式**：如果文档中有列表，识别项目符号或编号格式
- **页眉页脚**：如果文档中有页眉页脚，识别其格式

## ⚠️ 关键要求（必须严格遵守）：

1. **必须从文档中实际提取所有格式信息，绝对不要使用默认值或猜测**
2. **如果文档中某个格式信息不存在，不要猜测，应该明确说明或使用文档中实际存在的值**
3. **字体名称必须准确**：如果文档中是"宋体"，就返回"宋体"；如果是"楷体_GB2312"，就返回"楷体_GB2312"，不要替换
4. **字体大小必须准确**：如果文档中是12磅，就返回12；如果是16磅，就返回16，不要使用"通常"的值
5. **行距必须准确**：如果文档中是1.5倍，就返回1.5；如果是1.0倍，就返回1.0，不要猜测
6. **首行缩进必须准确**：如果文档中是24磅，就返回24；如果是0，就返回0，不要猜测
7. **标题格式必须准确**：从文档中实际提取标题的字体、大小、是否加粗，不要使用"通常"的格式

## 格式提取示例：

假设文档中提取到的格式信息显示：
- 正文段落字体：`font_name: "宋体"`, `font_size: "12"` → 返回 `"name": "宋体", "size": 12`
- 正文段落行距：`line_spacing: "1.5"` → 返回 `"line_spacing": 1.5`
- 正文首行缩进：`first_line_indent: "24"` → 返回 `"first_line_indent": 24`
- 一级标题字体：`font_name: "黑体"`, `font_size: "15"`, `bold: true` → 返回 `"font_name": "黑体", "font_size": 15, "bold": true`

**不要这样做**：
- ❌ 文档中是"楷体_GB2312"，但返回"宋体"（因为"通常"用宋体）
- ❌ 文档中是16磅，但返回12磅（因为"通常"是12磅）
- ❌ 文档中是1.0倍行距，但返回1.5倍（因为"通常"是1.5倍）
- ❌ 文档中没有首行缩进（0），但返回24磅（因为"通常"需要首行缩进）

## 输出格式要求（方案B：一次生成两个输出）：

**重要**：你需要同时生成两个输出：
1. **自然语言格式描述**（用于展示给用户）：用自然语言描述文档的格式要求，包括字体、字号、行距、页边距、标题格式等，让用户能够清晰理解格式要求。
2. **JSON格式指令**（用于系统执行）：严格按照JSON格式返回格式化指令，用于系统执行格式化操作。

**输出结构**：
```
【格式要求描述】

（这里用自然语言描述格式要求，包括但不限于：
- 正文格式：字体、字号、行距、首行缩进等
- 标题格式：各级标题的字体、字号、对齐方式等
- 页面设置：页边距、纸张大小等
- 特殊格式：目录、摘要、关键词等格式要求
- 布局规则：空行、间距等要求）

【格式化指令】

（这里返回JSON格式的格式化指令，严格按照下面的JSON结构）
```

**重要**：这是格式化指令，不是格式数据。指令描述了如何格式化文档，其中包含格式信息。**必须从文档中实际提取所有格式信息，不要使用默认值或猜测。如果文档中没有某个格式信息，不要猜测，应该使用文档中实际存在的值。**

格式化指令结构（推荐使用优化后的格式，更清晰易用）：
{{
  "version": "1.0",
  "description": "文档格式化指令描述（如：XX大学XX学位论文格式化指令）",
  "instruction_type": "format_application",
  "format_rules": {{
    "default_font": {{
      "name": "默认中文字体（必须从文档中提取，如：宋体、楷体_GB2312）",
      "size_pt": 字体大小（数字，单位：磅，必须从文档中提取，如：12）,
      "color": "颜色（RGB值，如：000000表示黑色，必须从文档中提取）"
    }},
    "english_font": {{
      "name": "英文字体（必须从文档中提取，如：Times New Roman、Arial）",
      "size_pt": 字体大小（数字，单位：磅，必须从文档中提取，如：12）
    }},
    "page": {{
      "size": "纸张大小（必须从文档中提取或根据页面尺寸推断，如：A4、Letter）",
      "margins": {{
        "top_cm": 上边距（数字，单位：厘米，必须从文档中提取，如：2.54）,
        "bottom_cm": 下边距（数字，单位：厘米，必须从文档中提取，如：2.54）,
        "left_cm": 左边距（数字，单位：厘米，必须从文档中提取，如：2.6）,
        "right_cm": 右边距（数字，单位：厘米，必须从文档中提取，如：2.6）
      }}
    }},
    "headings": {{
      "h1": {{
        "font_name": "字体名称（必须从文档中提取，不要使用默认值）",
        "font_size_pt": 字体大小（数字，单位：磅，必须从文档中提取。注意：如果文档中提到'小三号'、'15磅'或类似表述，必须使用15，不要使用默认的14）,
        "bold": true/false（必须从文档中提取）,
        "alignment": "对齐方式（必须从文档中提取，如：left、center）",
        "spacing_before_pt": 段前距（数字，单位：磅，必须从文档中提取，可选，如果提供了spacing_before_lines则此字段可选）,
        "spacing_after_pt": 段后距（数字，单位：磅，必须从文档中提取，可选，如果提供了spacing_after_lines则此字段可选）,
        "spacing_before_lines": 标题前空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段。如果只有磅数信息，可以转换为空行数）,
        "spacing_after_lines": 标题后空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段。如果只有磅数信息，可以转换为空行数）,
        "keep_with_next": true/false（可选，标题是否与下一段同页）
      }},
      "h2": {{
        "font_name": "字体名称（必须从文档中提取）",
        "font_size_pt": 字体大小（数字，单位：磅，必须从文档中提取）,
        "bold": true/false（必须从文档中提取）,
        "alignment": "对齐方式（必须从文档中提取）",
        "spacing_before_pt": 段前距（数字，单位：磅，必须从文档中提取，可选，如果提供了spacing_before_lines则此字段可选）,
        "spacing_after_pt": 段后距（数字，单位：磅，必须从文档中提取，可选，如果提供了spacing_after_lines则此字段可选）,
        "spacing_before_lines": 标题前空行数（数字，0-10，优先使用此字段）,
        "spacing_after_lines": 标题后空行数（数字，0-10，优先使用此字段）
      }},
      "h3": {{
        "font_name": "字体名称（必须从文档中提取）",
        "font_size_pt": 字体大小（数字，单位：磅，必须从文档中提取）,
        "bold": true/false（必须从文档中提取）,
        "alignment": "对齐方式（必须从文档中提取）",
        "spacing_before_pt": 段前距（数字，单位：磅，必须从文档中提取，可选，如果提供了spacing_before_lines则此字段可选）,
        "spacing_after_pt": 段后距（数字，单位：磅，必须从文档中提取，可选，如果提供了spacing_after_lines则此字段可选）,
        "spacing_before_lines": 标题前空行数（数字，0-10，优先使用此字段）,
        "spacing_after_lines": 标题后空行数（数字，0-10，优先使用此字段）
      }}
    }},
    "paragraph": {{
      "alignment": "对齐方式（left/center/right/justify，必须从文档中提取）",
      "line_spacing": 行距倍数（数字，如：1.5，必须从文档中提取）,
      "first_line_indent_chars": 首行缩进字符数（数字，如：2表示缩进2字符，必须从文档中提取）
    }},
    "special_sections": {{
      "title": {{
        "font_name": "论文标题字体（必须从文档中提取）",
        "font_size_pt": 字体大小（数字，单位：磅，必须从文档中提取）,
        "bold": true/false（必须从文档中提取）,
        "alignment": "对齐方式（通常为center，必须从文档中提取）",
        "spacing_after_pt": 标题后间距（数字，单位：磅，必须从文档中提取）
      }},
      "abstract": {{
        "label_text": "摘要标签文本（如：[摘要]、摘要等，必须从文档中提取）",
        "label_font": "标签字体（必须从文档中提取）",
        "label_size_pt": 标签字号（数字，单位：磅，必须从文档中提取）,
        "content_font": "摘要正文字体（必须从文档中提取）",
        "content_size_pt": 摘要正文字号（数字，单位：磅，必须从文档中提取）,
        "line_spacing": 行距倍数（数字，必须从文档中提取）,
        "word_count_min": 最少字数（数字，可选）,
        "word_count_max": 最多字数（数字，可选）
      }},
      "keywords": {{
        "label_text": "关键词标签文本（如：[关键词]、关键词等，必须从文档中提取）",
        "separator": "关键词分隔符（如：；、,等，必须从文档中提取）",
        "font": "关键词字体（必须从文档中提取）",
        "size_pt": 关键词字号（数字，单位：磅，必须从文档中提取）,
        "count_min": 最少关键词数（数字，可选）,
        "count_max": 最多关键词数（数字，可选）
      }},
      "abstract_english": {{
        "title_text": "英文摘要标题文本（如：Abstract等，必须从文档中提取）",
        "title_font": "英文摘要标题字体（必须从文档中提取，如：Times New Roman）",
        "title_size_pt": 英文摘要标题字号（数字，单位：磅，必须从文档中提取）,
        "title_bold": true/false（标题是否加粗，必须从文档中提取）,
        "content_font": "英文摘要正文字体（必须从文档中提取，如：Times New Roman）",
        "content_size_pt": 英文摘要正文字号（数字，单位：磅，必须从文档中提取）,
        "line_spacing": 行距倍数（数字，必须从文档中提取）
      }},
      "keywords_english": {{
        "label_text": "英文关键词标签文本（如：Key words等，必须从文档中提取）",
        "separator": "关键词分隔符（如：;、,等，必须从文档中提取）",
        "font": "英文关键词字体（必须从文档中提取，如：Times New Roman）",
        "size_pt": 英文关键词字号（数字，单位：磅，必须从文档中提取）,
        "label_bold": true/false（标签是否加粗，必须从文档中提取）
      }},
      "title_english": {{
        "font_name": "英文题目字体（必须从文档中提取，如：Times New Roman）",
        "font_size_pt": 英文题目字号（数字，单位：磅，必须从文档中提取）,
        "bold": true/false（是否加粗，必须从文档中提取）,
        "alignment": "对齐方式（通常为center，必须从文档中提取）"
      }},
      "author_info": {{
        "font_name": "作者信息字体（专业、年级、姓名等，必须从文档中提取）",
        "font_size_pt": 作者信息字号（数字，单位：磅，必须从文档中提取）,
        "alignment": "对齐方式（通常为center，必须从文档中提取）"
      }},
      "appendix": {{
        "title_text": "附录标题文本（如：附录、Appendix等，必须从文档中提取）",
        "title_font": "附录标题字体（必须从文档中提取）",
        "title_size_pt": 附录标题字号（数字，单位：磅，必须从文档中提取）,
        "title_alignment": "附录标题对齐方式（必须从文档中提取）",
        "content_font": "附录内容字体（必须从文档中提取）",
        "content_size_pt": 附录内容字号（数字，单位：磅，必须从文档中提取）
      }},
      "conclusion": {{
        "title_text": "结论标题文本（如：结　　论、结　论、结论等，必须从文档中提取。注意：如果文档中提到'结论'，优先使用'结　　论'（两个全角空格））",
        "title_font": "结论标题字体（必须从文档中提取）",
        "title_size_pt": 结论标题字号（数字，单位：磅，必须从文档中提取）,
        "title_alignment": "结论标题对齐方式（必须从文档中提取）",
        "title_spacing_before_pt": 标题前间距（数字，单位：磅，必须从文档中提取，可选，如果提供了title_before_lines则此字段可选）,
        "title_spacing_after_pt": 标题后间距（数字，单位：磅，必须从文档中提取，可选，如果提供了title_after_lines则此字段可选）,
        "title_before_lines": 标题前空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段）,
        "title_after_lines": 标题后空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段）
      }},
      "references": {{
        "title_text": "参考文献标题文本（如：参 考 文 献等，必须从文档中提取）",
        "title_font": "参考文献标题字体（必须从文档中提取）",
        "title_size_pt": 参考文献标题字号（数字，单位：磅，必须从文档中提取）,
        "title_alignment": "参考文献标题对齐方式（必须从文档中提取）",
        "title_spacing_before_pt": 标题前间距（数字，单位：磅，必须从文档中提取，可选，如果提供了title_before_lines则此字段可选）,
        "title_spacing_after_pt": 标题后间距（数字，单位：磅，必须从文档中提取，可选，如果提供了title_after_lines则此字段可选）,
        "title_before_lines": 标题前空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段）,
        "title_after_lines": 标题后空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段）,
        "item_font": "参考文献条目字体（必须从文档中提取）",
        "item_size_pt": 参考文献条目字号（数字，单位：磅，必须从文档中提取）,
        "line_spacing": 行距倍数（数字，必须从文档中提取）,
        "item_alignment": "参考文献条目对齐方式（通常为left，必须从文档中提取）",
        "format_rules": {{
          "journal_format": "期刊格式模板（如：[序号] 作者. 文题[J]. 刊名, 年, 卷号(期号): 起-止页码.，必须从文档中提取）",
          "book_format": "专著格式模板（如：[序号] 作者. 书名[M]. 出版地：出版者，出版年. 起-止页码.，必须从文档中提取）",
          "numbering_style": "序号样式（如：[1], [2]等，必须从文档中提取）"
        }}
      }},
      "acknowledgement": {{
        "title_text": "致谢标题文本（如：致　谢等，必须从文档中提取）",
        "title_font": "致谢标题字体（必须从文档中提取）",
        "title_size_pt": "致谢标题字号（数字，单位：磅，必须从文档中提取）",
        "title_alignment": "致谢标题对齐方式（必须从文档中提取）",
        "title_spacing_before_pt": 标题前间距（数字，单位：磅，必须从文档中提取，可选，如果提供了title_before_lines则此字段可选）,
        "title_spacing_after_pt": 标题后间距（数字，单位：磅，必须从文档中提取，可选，如果提供了title_after_lines则此字段可选）,
        "title_before_lines": 标题前空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段）,
        "title_after_lines": 标题后空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段）
      }},
      "table_of_contents": {{
        "title_text": "目录标题文本（如：目　　录、目　录、目 录等，必须从文档中提取。注意：如果文档中提到'目录'，优先使用'目　　录'（两个全角空格））",
        "title_font": "目录标题字体（必须从文档中提取）",
        "title_size_pt": 目录标题字号（数字，单位：磅，必须从文档中提取）,
        "title_alignment": "目录标题对齐方式（必须从文档中提取）",
        "title_spacing_before_pt": 标题前间距（数字，单位：磅，必须从文档中提取，可选，如果提供了title_before_lines则此字段可选）,
        "title_spacing_after_pt": 标题后间距（数字，单位：磅，必须从文档中提取，可选，如果提供了title_after_lines则此字段可选）,
        "title_before_lines": 标题前空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段）,
        "title_after_lines": 标题后空行数（数字，0-10，优先使用此字段。如果Word文档中明确提到"空X行"，直接使用此字段）,
        "entry_levels": [1, 2, 3]（目录支持的级别，必须从文档中提取）,
        "entry_format": {{
          "level_1": {{
            "font": "一级目录条目字体（必须从文档中提取）",
            "size_pt": 一级目录条目字号（数字，单位：磅，必须从文档中提取）,
            "bold": true/false（是否加粗，必须从文档中提取）
          }},
          "level_2": {{
            "font": "二级目录条目字体（必须从文档中提取）",
            "size_pt": 二级目录条目字号（数字，单位：磅，必须从文档中提取）,
            "bold": true/false（是否加粗，必须从文档中提取）
          }},
          "level_3": {{
            "font": "三级目录条目字体（必须从文档中提取）",
            "size_pt": 三级目录条目字号（数字，单位：磅，必须从文档中提取）,
            "bold": true/false（是否加粗，必须从文档中提取）
          }}
        }},
        "line_spacing": 目录行距倍数（数字，如：1.5，必须从文档中提取）
      }}
    }},
    "application_rules": {{
      "heading_detection": "标题识别规则（如：通过段落前缀数字识别层级）",
      "special_section_detection": [
        {{"marker": "结　论", "type": "conclusion"}},
        {{"marker": "参 考 文 献", "type": "references"}},
        {{"marker": "致　谢", "type": "acknowledgement"}},
        {{"marker": "附录", "type": "appendix"}},
        {{"marker": "Abstract", "type": "abstract_english"}},
        {{"marker": "Key words", "type": "keywords_english"}}
      ],
      "document_structure": {{
        "section_order": [
          "封面",
          "原创性声明",
          "评审表",
          "答辩记录表",
          "目录",
          "中文题目",
          "摘要",
          "关键词",
          "英文题目",
          "Abstract",
          "Key words",
          "正文",
          "结论",
          "参考文献",
          "附录",
          "致谢"
        ],
        "required_sections": ["目录", "摘要", "关键词", "正文", "结论", "参考文献"],
        "optional_sections": ["Abstract", "Key words", "附录", "致谢"]
      }},
      "auto_generate_toc": true/false（是否自动生成目录，如果模板中有目录章节则为false，否则为true）,
      "toc_generation_rules": {{
        "include_levels": [1, 2, 3]（包含的目录级别，必须从文档中提取）,
        "exclude_sections": ["摘要", "关键词", "目录"]（排除的章节，必须从文档中提取）,
        "page_number_format": "arabic"（页码格式：arabic/roman，必须从文档中提取）
      }},
      "font_fallback": {{
        "chinese": "中文字体回退（如：宋体）",
        "english": "英文字体回退（如：Times New Roman）"
      }},
      "chapter_numbering_format": {{
        "level_1": {{
          "format_type": "chinese_chapter/numbered/roman"（格式类型，必须从文档中提取）,
          "pattern": "格式模板（如：第{{number}}章 {{title}}、{{number}}.{{sub}} {{title}}等，必须从文档中提取）",
          "number_style": "chinese/arabic/roman"（数字样式，必须从文档中提取）,
          "examples": ["第一章 引言", "1.1 研究背景"]（格式示例，必须从文档中提取）
        }},
        "level_2": {{
          "format_type": "numbered"（格式类型，必须从文档中提取）,
          "pattern": "格式模板（如：{{parent}}.{{number}} {{title}}，必须从文档中提取）",
          "number_style": "arabic"（数字样式，必须从文档中提取）,
          "examples": ["1.1 研究背景", "1.2 研究意义"]（格式示例，必须从文档中提取）
        }},
        "level_3": {{
          "format_type": "numbered"（格式类型，必须从文档中提取）,
          "pattern": "格式模板（如：{{parent}}.{{sub}}.{{number}} {{title}}，必须从文档中提取）",
          "number_style": "arabic"（数字样式，必须从文档中提取）,
          "examples": ["1.1.1 人工智能定义", "1.1.2 发展历程"]（格式示例，必须从文档中提取）
        }}
      }},
      "special_section_format_rules": {{
        "abstract": {{
          "title": "摘要"（标题文本，必须从文档中提取）,
          "should_have_numbering": false（是否应该有章节编号，必须从文档中提取）,
          "position": "after_toc"（位置：after_toc/after_title/before_body，必须从文档中提取）
        }},
        "keywords": {{
          "title": "关键词"（标题文本，必须从文档中提取）,
          "should_have_numbering": false（是否应该有章节编号，必须从文档中提取）,
          "position": "after_abstract"（位置，必须从文档中提取）
        }},
        "conclusion": {{
          "title": "结语/结论"（标题文本，必须从文档中提取，注意：有些学校用"结语"，有些用"结论"）,
          "should_have_numbering": false（是否应该有章节编号，必须从文档中提取）,
          "position": "before_references"（位置，必须从文档中提取）
        }},
        "references": {{
          "title": "参考文献"（标题文本，必须从文档中提取）,
          "should_have_numbering": false（是否应该有章节编号，必须从文档中提取）,
          "position": "before_appendix"（位置，必须从文档中提取）
        }},
        "acknowledgement": {{
          "title": "致谢"（标题文本，必须从文档中提取）,
          "should_have_numbering": false（是否应该有章节编号，必须从文档中提取）,
          "position": "last"（位置，必须从文档中提取）
        }}
      }}
    }}
  }}
}}

## 输出要求：
1. 先输出自然语言格式描述（【格式要求描述】部分）
2. 再输出JSON格式指令（【格式化指令】部分，用```json代码块包裹）
3. 所有格式必须从文档中实际提取，不使用默认值
4. JSON必须有效，数字为数字类型，字体名称准确
5. 目录格式需特别注意：提取标题前后空行、条目格式（多级）、缩进值"""
        
        return prompt
    
    @classmethod
    async def format_thesis(
        cls,
        query_db: AsyncSession,
        thesis_id: int,
        word_file_path: Optional[str] = None,
        format_instructions: Optional[str] = None,
        config_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        格式化论文（根据AI提取的格式指令）
        
        :param query_db: 数据库会话
        :param thesis_id: 论文ID
        :param word_file_path: Word文档路径（如果提供，会先读取并提取格式指令）
        :param format_instructions: 格式化指令（JSON字符串，如果提供则直接使用）
        :param config_id: AI模型配置ID（可选）
        :return: 格式化结果
        """
        if not DOCX_AVAILABLE:
            import sys
            python_path = sys.executable
            error_msg = (
                f'python-docx 未安装，无法格式化Word文档。\n'
                f'当前Python路径: {python_path}\n'
                f'请执行以下命令安装: pip install python-docx\n'
                f'如果使用虚拟环境，请确保已激活虚拟环境。'
            )
            logger.error(error_msg)
            raise ServiceException(message=error_msg)
        
        try:
            # 如果没有提供格式指令，需要先读取Word文档并提取
            if not format_instructions and word_file_path:
                logger.info(f"开始读取Word文档并提取格式指令 - 文件: {word_file_path}")
                read_result = await cls.read_word_document_with_ai(query_db, word_file_path, config_id)
                format_instructions = read_result['format_instructions']
            
            if not format_instructions:
                raise ServiceException(message='未提供格式化指令，无法进行格式化')
            
            # 解析格式化指令
            logger.info(f"[格式化流程] 解析格式化指令")
            logger.debug(f"  格式指令长度: {len(format_instructions) if format_instructions else 0} 字符")
            try:
                format_instruction_data = json.loads(format_instructions) if isinstance(format_instructions, str) else format_instructions
                logger.info(f"  成功解析JSON格式指令")
            except json.JSONDecodeError:
                # 如果AI返回的不是纯JSON，尝试提取JSON部分
                logger.warning(f"  JSON解析失败，尝试提取JSON部分")
                format_instruction_data = cls._extract_json_from_text(format_instructions)
            
            # 从格式化指令中提取格式配置和布局规则
            # 格式化指令可能包含 format_rules 和 layout_rules 字段，或者直接就是格式配置
            layout_rules = {}
            if isinstance(format_instruction_data, dict):
                if 'format_rules' in format_instruction_data:
                    # 新格式：格式化指令包含 format_rules
                    format_config = format_instruction_data['format_rules']
                    layout_rules = format_instruction_data.get('layout_rules', {})
                    logger.info(f"  使用新格式指令（包含format_rules）")
                    
                    # 检查是否是优化后的新格式（包含 default_font, english_font 等）
                    if 'default_font' in format_config or 'special_sections' in format_config:
                        logger.info(f"  检测到优化后的新格式，转换为兼容格式")
                        format_config, extracted_layout_rules = cls._convert_optimized_format_to_legacy(format_config)
                        # 合并提取的布局规则
                        if extracted_layout_rules:
                            if not layout_rules:
                                layout_rules = {}
                            if 'section_spacing' not in layout_rules:
                                layout_rules['section_spacing'] = {}
                            layout_rules['section_spacing'].update(extracted_layout_rules.get('section_spacing', {}))
                else:
                    # 旧格式：直接就是格式配置
                    format_config = format_instruction_data
                    logger.info(f"  使用旧格式指令（直接是格式配置）")
            else:
                format_config = format_instruction_data
                logger.warning(f"  格式指令不是字典类型，使用原值")
            
            # 验证和修正格式数据（确保格式符合标准要求）
            logger.info(f"[格式化流程] 验证和修正格式配置")
            format_config = cls._validate_and_fix_format_config(format_config)
            logger.info(f"  格式配置验证完成")
            
            # 获取论文的所有章节（只包含已完成的章节）
            from module_thesis.dao.thesis_dao import ThesisChapterDao
            logger.info(f"[格式化流程] 开始格式化论文 - 论文ID: {thesis_id}")
            all_chapters = await ThesisChapterDao.get_chapter_list_by_thesis(query_db, thesis_id)
            logger.info(f"  获取到 {len(all_chapters)} 个章节（所有状态）")
            
            # 只格式化已完成的章节
            chapters = [c for c in all_chapters if c.status == 'completed']
            logger.info(f"  已完成章节数量: {len(chapters)}")
            
            if not chapters:
                logger.error(f"  没有已完成的章节，无法格式化")
                raise ServiceException(message='论文没有已完成的章节内容，无法格式化')
            
            # 列出所有章节标题
            chapter_titles = [c.title for c in chapters]
            logger.info(f"  章节列表: {', '.join(chapter_titles[:5])}{'...' if len(chapter_titles) > 5 else ''}")
            
            # 注意：不再进行格式转换
            # 因为大纲和章节生成时已经根据格式指令使用了正确的格式
            # 这里只需要直接应用格式即可
            logger.info(f"[格式化流程] 跳过格式转换（内容已符合格式要求）")
            
            # 如果仍然存在旧的转换规则配置，记录警告但不执行
            application_rules = format_config.get('application_rules', {})
            if application_rules.get('chapter_numbering_conversion', {}).get('enabled', False):
                logger.warning(f"  检测到旧的转换规则配置，但已跳过（内容应在生成时已符合格式要求）")
            if application_rules.get('abstract_extraction', {}).get('enabled', False):
                logger.warning(f"  检测到旧的提取规则配置，但已跳过（内容应在生成时已符合格式要求）")
            
            # 获取论文基本信息（用于文档头部）
            from module_thesis.dao.thesis_dao import ThesisDao
            thesis = await ThesisDao.get_thesis_by_id(query_db, thesis_id)
            logger.info(f"  论文标题: {thesis.title if thesis and thesis.title else 'N/A'}")
            
            # 创建格式化的Word文档（传入格式配置和布局规则）
            logger.info(f"  开始创建格式化文档...")
            output_path = cls._create_formatted_document(chapters, format_config, thesis_id, thesis, layout_rules)
            
            logger.info(f"[格式化流程] 论文格式化完成 - 论文ID: {thesis_id}, 输出文件: {output_path}")
            
            return {
                'formatted_file_path': output_path,
                'format_instructions': format_instructions,  # 原始格式化指令
                'format_config': format_config  # 从指令中提取的格式配置
            }
            
        except ServiceException:
            raise
        except Exception as e:
            logger.error(f"格式化论文失败: {str(e)}", exc_info=True)
            raise ServiceException(message=f'格式化论文失败: {str(e)}')
    
    @classmethod
    def _convert_optimized_format_to_legacy(cls, optimized_format: Dict[str, Any]) -> tuple[Dict[str, Any], Dict[str, Any]]:
        """
        将优化后的新格式转换为旧格式（保持向后兼容）
        
        :param optimized_format: 优化后的格式配置
        :return: (转换后的旧格式配置, 提取的布局规则)
        """
        legacy_format = {}
        extracted_layout_rules = {}
        
        # 转换默认字体
        if 'default_font' in optimized_format:
            default_font = optimized_format['default_font']
            legacy_format['font'] = {
                'name': default_font.get('name', '宋体'),
                'size': default_font.get('size_pt', 12),
                'color': default_font.get('color', '000000')
            }
        elif 'font' in optimized_format:
            legacy_format['font'] = optimized_format['font']
        
        # 转换段落格式
        if 'paragraph' in optimized_format:
            para = optimized_format['paragraph'].copy()
            # 转换首行缩进：字符数转磅数（1字符≈12磅）
            if 'first_line_indent_chars' in para:
                indent_chars = para.pop('first_line_indent_chars')
                para['first_line_indent'] = indent_chars * 12
            legacy_format['paragraph'] = para
        
        # 转换标题格式
        if 'headings' in optimized_format:
            headings = {}
            for level in ['h1', 'h2', 'h3']:
                if level in optimized_format['headings']:
                    heading = optimized_format['headings'][level].copy()
                    # 转换字段名
                    if 'font_size_pt' in heading:
                        heading['font_size'] = heading.pop('font_size_pt')
                    
                    # 优先使用空行数字段，如果没有则从磅数字段转换
                    spacing_before_lines = heading.get('spacing_before_lines', None)
                    spacing_after_lines = heading.get('spacing_after_lines', None)
                    
                    if spacing_before_lines is None:
                        # 如果没有空行数字段，则从磅数字段转换
                        spacing_before_pt = heading.get('spacing_before_pt', 0)
                        # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                        if spacing_before_pt >= 24:
                            spacing_before_lines = int(spacing_before_pt / 24)
                        elif spacing_before_pt > 0:
                            spacing_before_lines = 1
                        else:
                            spacing_before_lines = 0
                    else:
                        spacing_before_lines = int(spacing_before_lines) if spacing_before_lines is not None else 0
                    
                    if spacing_after_lines is None:
                        # 如果没有空行数字段，则从磅数字段转换
                        spacing_after_pt = heading.get('spacing_after_pt', 0)
                        # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                        if spacing_after_pt >= 24:
                            spacing_after_lines = int(spacing_after_pt / 24)
                        elif spacing_after_pt > 0:
                            spacing_after_lines = 1
                        else:
                            spacing_after_lines = 0
                    else:
                        spacing_after_lines = int(spacing_after_lines) if spacing_after_lines is not None else 0
                    
                    # 保留磅数字段（向后兼容），但优先使用空行数
                    if 'spacing_before_pt' in heading:
                        heading['spacing_before'] = heading.pop('spacing_before_pt')
                    if 'spacing_after_pt' in heading:
                        heading['spacing_after'] = heading.pop('spacing_after_pt')
                    
                    # 添加空行数字段到布局规则
                    if 'section_spacing' not in extracted_layout_rules:
                        extracted_layout_rules['section_spacing'] = {}
                    if 'headings' not in extracted_layout_rules['section_spacing']:
                        extracted_layout_rules['section_spacing']['headings'] = {}
                    extracted_layout_rules['section_spacing']['headings'][level] = {
                        'before': spacing_before_lines,
                        'after': spacing_after_lines
                    }
                    
                    headings[level] = heading
            legacy_format['headings'] = headings
        
        # 转换页面设置
        if 'page' in optimized_format:
            page = optimized_format['page'].copy()
            if 'margins' in page:
                margins = page['margins']
                # 转换厘米到磅（1厘米 = 28.35磅）
                legacy_margins = {}
                if 'top_cm' in margins:
                    legacy_margins['top'] = margins['top_cm'] * 28.35
                if 'bottom_cm' in margins:
                    legacy_margins['bottom'] = margins['bottom_cm'] * 28.35
                if 'left_cm' in margins:
                    legacy_margins['left'] = margins['left_cm'] * 28.35
                if 'right_cm' in margins:
                    legacy_margins['right'] = margins['right_cm'] * 28.35
                page['margins'] = legacy_margins
            legacy_format['page'] = page
        
        # 转换特殊格式（special_sections -> special_formats）
        if 'special_sections' in optimized_format:
            special_sections = optimized_format['special_sections']
            special_formats = {}
            
            # 转换摘要
            if 'abstract' in special_sections:
                abstract = special_sections['abstract']
                special_formats['abstract'] = {
                    'title_format': {
                        'font_name': abstract.get('label_font', '宋体'),
                        'font_size': abstract.get('label_size_pt', 12),
                        'bold': False,
                        'alignment': 'left'
                    },
                    'content_format': {
                        'font_name': abstract.get('content_font', '宋体'),
                        'font_size': abstract.get('content_size_pt', 12),
                        'line_spacing': abstract.get('line_spacing', 1.5)
                    }
                }
            
            # 转换关键词
            if 'keywords' in special_sections:
                keywords = special_sections['keywords']
                special_formats['keywords'] = {
                    'font_name': keywords.get('font', '宋体'),
                    'font_size': keywords.get('size_pt', 12),
                    'line_spacing': 1.5
                }
            
            # 转换结论
            if 'conclusion' in special_sections:
                conclusion = special_sections['conclusion']
                special_formats['conclusion'] = {
                    'title_format': {
                        'font_name': conclusion.get('title_font', '黑体'),
                        'font_size': conclusion.get('title_size_pt', 15),
                        'bold': True,
                        'alignment': conclusion.get('title_alignment', 'center')
                    },
                    'content_format': {
                        'font_name': '宋体',
                        'font_size': 12,
                        'line_spacing': 1.5
                    }
                }
                
                # 提取结论标题前后空行设置
                title_before_lines = conclusion.get('title_before_lines', None)
                title_after_lines = conclusion.get('title_after_lines', None)
                
                if title_before_lines is None:
                    title_spacing_before_pt = conclusion.get('title_spacing_before_pt', 0)
                    # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                    if title_spacing_before_pt >= 24:
                        title_before_lines = int(title_spacing_before_pt / 24)
                    elif title_spacing_before_pt > 0:
                        title_before_lines = 1
                    else:
                        title_before_lines = 0
                else:
                    title_before_lines = int(title_before_lines) if title_before_lines is not None else 0
                
                if title_after_lines is None:
                    title_spacing_after_pt = conclusion.get('title_spacing_after_pt', 0)
                    # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                    if title_spacing_after_pt >= 24:
                        title_after_lines = int(title_spacing_after_pt / 24)
                    elif title_spacing_after_pt > 0:
                        title_after_lines = 1
                    else:
                        title_after_lines = 0
                else:
                    title_after_lines = int(title_after_lines) if title_after_lines is not None else 0
                
                if 'section_spacing' not in extracted_layout_rules:
                    extracted_layout_rules['section_spacing'] = {}
                extracted_layout_rules['section_spacing']['conclusion'] = {
                    'title_before': title_before_lines,
                    'title_after': title_after_lines
                }
                logger.info(f"  提取结论标题空行设置: 前{title_before_lines}行, 后{title_after_lines}行")
            
            # 转换目录
            if 'table_of_contents' in special_sections:
                toc = special_sections['table_of_contents']
                special_formats['table_of_contents'] = {
                    'title_format': {
                        'font_name': toc.get('title_font', '黑体'),
                        'font_size': toc.get('title_size_pt', 15),
                        'bold': True,
                        'alignment': toc.get('title_alignment', 'center')
                    },
                    'entry_format': {
                        'font_name': '宋体',
                        'font_size': 12,
                        'line_spacing': 1.5,
                        'alignment': 'justify'
                    }
                }
                
                # 提取目录标题前后空行设置（转换为布局规则）
                # 优先使用空行数字段（title_before_lines、title_after_lines）
                # 如果没有空行数字段，则从磅数字段转换（保持向后兼容）
                title_before_lines = toc.get('title_before_lines', None)
                title_after_lines = toc.get('title_after_lines', None)
                
                # 如果没有空行数字段，则从磅数字段转换
                if title_before_lines is None:
                    title_spacing_before_pt = toc.get('title_spacing_before_pt', 0)
                    # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                    if title_spacing_before_pt >= 24:
                        title_before_lines = int(title_spacing_before_pt / 24)
                    elif title_spacing_before_pt > 0:
                        title_before_lines = 1
                    else:
                        title_before_lines = 0
                else:
                    title_before_lines = int(title_before_lines) if title_before_lines is not None else 0
                
                if title_after_lines is None:
                    title_spacing_after_pt = toc.get('title_spacing_after_pt', 0)
                    # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                    if title_spacing_after_pt >= 24:
                        title_after_lines = int(title_spacing_after_pt / 24)
                    elif title_spacing_after_pt > 0:
                        title_after_lines = 1
                    else:
                        title_after_lines = 0
                else:
                    title_after_lines = int(title_after_lines) if title_after_lines is not None else 0
                
                # 添加到布局规则
                if 'section_spacing' not in extracted_layout_rules:
                    extracted_layout_rules['section_spacing'] = {}
                extracted_layout_rules['section_spacing']['table_of_contents'] = {
                    'title_before': title_before_lines,
                    'title_after': title_after_lines
                }
                logger.info(f"  提取目录标题空行设置: 前{title_before_lines}行, 后{title_after_lines}行")
            
            # 转换参考文献
            if 'references' in special_sections:
                refs = special_sections['references']
                special_formats['references'] = {
                    'title_format': {
                        'font_name': refs.get('title_font', '黑体'),
                        'font_size': refs.get('title_size_pt', 15),
                        'bold': True,
                        'alignment': refs.get('title_alignment', 'center')
                    },
                    'content_format': {
                        'font_name': refs.get('item_font', '宋体'),
                        'font_size': refs.get('item_size_pt', 10.5),
                        'line_spacing': refs.get('line_spacing', 1.5)
                    }
                }
                
                # 提取参考文献标题前后空行设置
                title_before_lines = refs.get('title_before_lines', None)
                title_after_lines = refs.get('title_after_lines', None)
                
                if title_before_lines is None:
                    title_spacing_before_pt = refs.get('title_spacing_before_pt', 0)
                    # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                    if title_spacing_before_pt >= 24:
                        title_before_lines = int(title_spacing_before_pt / 24)
                    elif title_spacing_before_pt > 0:
                        title_before_lines = 1
                    else:
                        title_before_lines = 0
                else:
                    title_before_lines = int(title_before_lines) if title_before_lines is not None else 0
                
                if title_after_lines is None:
                    title_spacing_after_pt = refs.get('title_spacing_after_pt', 0)
                    # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                    if title_spacing_after_pt >= 24:
                        title_after_lines = int(title_spacing_after_pt / 24)
                    elif title_spacing_after_pt > 0:
                        title_after_lines = 1
                    else:
                        title_after_lines = 0
                else:
                    title_after_lines = int(title_after_lines) if title_after_lines is not None else 0
                
                if 'section_spacing' not in extracted_layout_rules:
                    extracted_layout_rules['section_spacing'] = {}
                extracted_layout_rules['section_spacing']['references'] = {
                    'title_before': title_before_lines,
                    'title_after': title_after_lines
                }
                logger.info(f"  提取参考文献标题空行设置: 前{title_before_lines}行, 后{title_after_lines}行")
            
            # 转换致谢
            if 'acknowledgement' in special_sections:
                ack = special_sections['acknowledgement']
                special_formats['acknowledgement'] = {
                    'title_format': {
                        'font_name': ack.get('title_font', '黑体'),
                        'font_size': ack.get('title_size_pt', 15),
                        'bold': True,
                        'alignment': ack.get('title_alignment', 'center')
                    },
                    'content_format': {
                        'font_name': '宋体',
                        'font_size': 12,
                        'line_spacing': 1.5
                    }
                }
                
                # 提取致谢标题前后空行设置
                title_before_lines = ack.get('title_before_lines', None)
                title_after_lines = ack.get('title_after_lines', None)
                
                if title_before_lines is None:
                    title_spacing_before_pt = ack.get('title_spacing_before_pt', 0)
                    # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                    if title_spacing_before_pt >= 24:
                        title_before_lines = int(title_spacing_before_pt / 24)
                    elif title_spacing_before_pt > 0:
                        title_before_lines = 1
                    else:
                        title_before_lines = 0
                else:
                    title_before_lines = int(title_before_lines) if title_before_lines is not None else 0
                
                if title_after_lines is None:
                    title_spacing_after_pt = ack.get('title_spacing_after_pt', 0)
                    # 转换逻辑：24磅=1行，小于24磅但大于0视为1行
                    if title_spacing_after_pt >= 24:
                        title_after_lines = int(title_spacing_after_pt / 24)
                    elif title_spacing_after_pt > 0:
                        title_after_lines = 1
                    else:
                        title_after_lines = 0
                else:
                    title_after_lines = int(title_after_lines) if title_after_lines is not None else 0
                
                if 'section_spacing' not in extracted_layout_rules:
                    extracted_layout_rules['section_spacing'] = {}
                extracted_layout_rules['section_spacing']['acknowledgement'] = {
                    'title_before': title_before_lines,
                    'title_after': title_after_lines
                }
                logger.info(f"  提取致谢标题空行设置: 前{title_before_lines}行, 后{title_after_lines}行")
            
            if special_formats:
                legacy_format['special_formats'] = special_formats
        
        # 保留其他字段（application_rules, layout_rules等）
        for key in ['application_rules', 'layout_rules']:
            if key in optimized_format:
                legacy_format[key] = optimized_format[key]
        
        logger.info(f"已将优化格式转换为旧格式，保留字段: {list(legacy_format.keys())}")
        logger.info(f"提取的布局规则: {list(extracted_layout_rules.keys())}")
        return legacy_format, extracted_layout_rules
    
    @classmethod
    def _extract_json_from_text(cls, text: str) -> Dict[str, Any]:
        """从文本中提取JSON内容"""
        import re
        
        # 尝试提取JSON代码块
        json_match = re.search(r'```json\s*(\{.*?\})\s*```', text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group(1))
        
        # 尝试提取大括号内容
        json_match = re.search(r'(\{.*\})', text, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group(1))
            except:
                pass
        
        # 如果都失败，返回默认配置
        logger.warning("无法从AI响应中提取JSON，使用默认格式配置")
        return {
            'font': {'name': '宋体', 'size': 12},
            'paragraph': {'alignment': 'left', 'line_spacing': 1.5},
            'page': {'margins': {'top': 72, 'bottom': 72, 'left': 90, 'right': 90}}
        }
    
    @classmethod
    def _validate_and_fix_format_config(cls, format_config: Dict[str, Any]) -> Dict[str, Any]:
        """
        验证和修正格式配置，只修正明显错误，保留不同学校的格式差异
        
        注意：不同学校可能有不同的格式要求，因此只修正明显错误（如负数、极端值），
        保留从模板中提取的实际格式，不要强制修正为标准格式。
        
        :param format_config: 原始格式配置
        :return: 修正后的格式配置
        """
        config = format_config.copy()
        
        # 1. 验证和修正默认字体（只修正明显错误）
        if 'font' not in config:
            config['font'] = {}
        
        font_config = config['font']
        font_name = font_config.get('name', '宋体')
        # 只修正明显错误：如果字体名称是标题字体（黑体、楷体等），可能是误识别
        # 但保留其他合理的字体（如Times New Roman、Arial等）
        if '黑体' in font_name or '楷体' in font_name or 'Bold' in font_name:
            # 检查是否是明显的误识别（如默认字体是黑体）
            logger.warning(f"默认字体可能是误识别（{font_name}），但保留原值以支持不同学校格式")
            # 不强制修正，保留原值
        
        # 动态修正异常值：字体大小小于8磅或大于30磅
        font_size = font_config.get('size') or font_config.get('size_pt')
        if font_size:
            if isinstance(font_size, str):
                try:
                    font_size = float(font_size)
                except:
                    font_size = 12
            # 更严格的限制：8-30磅为合理范围
            if font_size < 8 or font_size > 30:
                logger.warning(f"默认字体大小异常值（{font_size}磅），动态修正为12磅")
                font_config['size'] = 12
                font_config['size_pt'] = 12
            else:
                # 确保两个字段都设置
                font_config['size'] = font_size
                font_config['size_pt'] = font_size
        else:
            # 如果没有设置，使用默认值
            font_config['size'] = 12
            font_config['size_pt'] = 12
        # 保留8-30磅之间的合理差异（不同学校可能有不同要求）
        
        # 2. 验证和修正段落格式（只修正明显错误）
        if 'paragraph' not in config:
            config['paragraph'] = {}
        
        para_config = config['paragraph']
        # 只修正极端值：行距小于0.5或大于3.0
        line_spacing = para_config.get('line_spacing', 1.5)
        if isinstance(line_spacing, str):
            try:
                line_spacing = float(line_spacing)
            except:
                line_spacing = 1.5
        if line_spacing < 0.5 or line_spacing > 3.0:
            logger.warning(f"行距极端值（{line_spacing}），修正为1.5")
            para_config['line_spacing'] = 1.5
        # 保留0.8-2.5之间的合理差异
        
        # 只修正负数：首行缩进为负数时修正为0（有些格式可能不需要首行缩进，保留0值）
        first_line_indent = para_config.get('first_line_indent', 24)
        if isinstance(first_line_indent, str):
            try:
                first_line_indent = float(first_line_indent)
            except:
                first_line_indent = 24
        if first_line_indent < 0:
            logger.warning(f"首行缩进为负数（{first_line_indent}磅），修正为0")
            para_config['first_line_indent'] = 0
        # 保留0或正数（不同学校可能有不同要求，有些不需要首行缩进）
        
        # 如果右缩进为负数，修正为0
        right_indent = para_config.get('right_indent', 0)
        if isinstance(right_indent, str):
            try:
                right_indent = float(right_indent)
            except:
                right_indent = 0
        if right_indent < 0:
            logger.warning(f"右缩进为负数（{right_indent}磅），修正为0")
            para_config['right_indent'] = 0
        
        # 3. 验证和修正标题格式（只修正明显错误，保留不同学校的格式差异）
        if 'headings' not in config:
            config['headings'] = {}
        
        headings_config = config['headings']
        
        # 标准标题格式（仅作为默认值，如果缺少配置时使用）
        default_headings = {
            'h1': {'font_name': '黑体', 'font_size': 15, 'bold': True, 'alignment': 'left'},
            'h2': {'font_name': '黑体', 'font_size': 14, 'bold': True, 'alignment': 'left'},
            'h3': {'font_name': '黑体', 'font_size': 12, 'bold': True, 'alignment': 'left'}
        }
        
        for level in ['h1', 'h2', 'h3']:
            if level in headings_config:
                heading = headings_config[level]
                
                # 不强制修正字体名称，保留从模板中提取的实际字体
                # 不同学校可能使用不同的标题字体（黑体、楷体、宋体加粗等）
                
                # 修正异常值：字体大小小于8磅或大于30磅（更严格的限制）
                font_size = heading.get('font_size', default_headings[level]['font_size'])
                if isinstance(font_size, str):
                    try:
                        font_size = float(font_size)
                    except:
                        font_size = default_headings[level]['font_size']
                # 更严格的限制：8-30磅为合理范围
                if font_size < 8 or font_size > 30:
                    logger.warning(f"{level}标题字体大小异常值（{font_size}磅），修正为{default_headings[level]['font_size']}磅")
                    heading['font_size'] = default_headings[level]['font_size']
                # 如果使用font_size_pt字段，也检查
                if 'font_size_pt' in heading:
                    font_size_pt = heading.get('font_size_pt')
                    if isinstance(font_size_pt, (int, float)):
                        if font_size_pt < 8 or font_size_pt > 30:
                            logger.warning(f"{level}标题字体大小异常值（font_size_pt: {font_size_pt}磅），修正为{default_headings[level]['font_size']}磅")
                            heading['font_size_pt'] = default_headings[level]['font_size']
                # 保留8-30磅之间的合理差异
                
                # 不强制修正加粗状态，保留从模板中提取的实际状态
                # 有些学校可能标题不加粗，或使用其他方式表示强调
            else:
                # 如果缺少某个级别的标题格式，添加默认格式（但不强制已有格式）
                headings_config[level] = default_headings[level].copy()
        
        # 4. 验证和修正特殊章节的字体大小（动态修正异常值）
        if 'special_sections' in config:
            special_sections = config['special_sections']
            
            # 检查所有特殊章节的字体大小
            sections_to_check = [
                ('abstract', ['label_size_pt', 'content_size_pt']),
                ('keywords', ['size_pt']),
                ('conclusion', ['title_size_pt']),
                ('references', ['title_size_pt', 'item_size_pt']),
                ('table_of_contents', ['title_size_pt']),
                ('title_english', ['font_size_pt']),
                ('author_info', ['font_size_pt']),
            ]
            
            for section_key, size_fields in sections_to_check:
                if section_key in special_sections:
                    section = special_sections[section_key]
                    for size_field in size_fields:
                        if size_field in section:
                            size = section[size_field]
                            if isinstance(size, (int, float)):
                                if size < 8 or size > 30:
                                    default_size = 12 if 'content' in size_field or 'item' in size_field else 14
                                    logger.warning(f"{section_key}.{size_field}字体大小异常值（{size}磅），动态修正为{default_size}磅")
                                    section[size_field] = default_size
        
        logger.info("格式配置验证完成（动态修正异常值，保留不同学校的格式差异）")
        return config
    
    @classmethod
    async def _get_universal_instruction_system(cls, query_db: AsyncSession) -> Dict[str, Any]:
        """
        从数据库读取完整指令系统
        
        :param query_db: 数据库会话
        :return: 完整指令系统（字典格式）
        """
        print("[获取完整指令系统] 开始从数据库读取...")
        logger.info("[获取完整指令系统] 开始从数据库读取...")
        try:
            instruction_system = await UniversalInstructionSystemDao.get_active_instruction_system(query_db)
            
            if not instruction_system:
                print("[获取完整指令系统] ⚠ 未找到激活的完整指令系统")
                logger.warning("未找到激活的完整指令系统，使用空字典")
                return {}
            
            print(f"[获取完整指令系统] ✓ 找到激活的指令系统")
            print(f"  ID: {instruction_system.id}")
            print(f"  版本: {instruction_system.version}")
            print(f"  描述: {instruction_system.description}")
            print(f"  是否激活: {instruction_system.is_active}")
            logger.info(f"[获取完整指令系统] ✓ 找到激活的指令系统 - ID: {instruction_system.id}, 版本: {instruction_system.version}")
            
            # 返回instruction_data字段（JSON格式）
            instruction_data = instruction_system.instruction_data
            data_type = type(instruction_data).__name__
            print(f"[获取完整指令系统] 指令数据类型: {data_type}")
            logger.info(f"[获取完整指令系统] 指令数据类型: {data_type}")
            
            if isinstance(instruction_data, dict):
                data_size = len(json.dumps(instruction_data, ensure_ascii=False))
                print(f"[获取完整指令系统] ✓ 指令数据为字典格式，大小: {data_size} 字符")
                logger.info(f"[获取完整指令系统] ✓ 指令数据为字典格式，大小: {data_size} 字符")
                return instruction_data
            elif isinstance(instruction_data, str):
                data_size = len(instruction_data)
                print(f"[获取完整指令系统] 指令数据为字符串格式，大小: {data_size} 字符，开始解析JSON...")
                logger.info(f"[获取完整指令系统] 指令数据为字符串格式，大小: {data_size} 字符，开始解析JSON...")
                parsed_data = json.loads(instruction_data)
                print(f"[获取完整指令系统] ✓ JSON解析成功")
                logger.info(f"[获取完整指令系统] ✓ JSON解析成功")
                return parsed_data
            else:
                print(f"[获取完整指令系统] ⚠ 指令数据格式异常: {data_type}")
                logger.warning(f"完整指令系统数据格式异常: {data_type}")
                return {}
                
        except json.JSONDecodeError as e:
            print(f"[获取完整指令系统] ✗ JSON解析失败: {str(e)}")
            logger.error(f"完整指令系统JSON解析失败: {str(e)}", exc_info=True)
            return {}
        except Exception as e:
            print(f"[获取完整指令系统] ✗ 读取失败: {str(e)}")
            logger.error(f"读取完整指令系统失败: {str(e)}", exc_info=True)
            return {}
    
    @classmethod
    async def _generate_format_instructions_directly(
        cls,
        query_db: AsyncSession,
        document_text: str,
        universal_instruction_system: Dict[str, Any],
        config_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        AI直接分析文档文本，一步生成自然语言描述和子集指令系统
        
        :param query_db: 数据库会话
        :param document_text: 文档文本内容
        :param universal_instruction_system: 完整指令系统
        :param config_id: AI模型配置ID
        :return: 包含自然语言描述和子集指令的字典
        """
        # 将完整指令系统转换为JSON字符串
        universal_system_json = json.dumps(universal_instruction_system, ensure_ascii=False, indent=2)
        if len(universal_system_json) > 10000:
            # 如果太长，只保留关键结构
            simplified_system = {
                'version': universal_instruction_system.get('version'),
                'description': universal_instruction_system.get('description'),
                'instruction_type': universal_instruction_system.get('instruction_type'),
                'format_rules': universal_instruction_system.get('format_rules', {}),
                'application_rules': universal_instruction_system.get('application_rules', {})
            }
            universal_system_json = json.dumps(simplified_system, ensure_ascii=False, indent=2)
        
        # 限制文档文本长度（避免超过AI上下文限制）
        # 重要：不能丢失任何文字信息！如果必须截断，要确保截断位置合理
        original_length = len(document_text)
        if original_length > 8000:
            logger.warning(f"文档文本过长（{original_length}字符），需要截断，但必须确保不丢失关键信息")
            print(f"[AI格式分析] ⚠ 警告：文档文本过长（{original_length}字符），需要截断")
            
            # 检查截断部分是否包含重要关键词
            truncated_part = document_text[8000:]
            important_keywords = ["空两行", "空行", "格式要求", "标题", "正文", "页边距", "字体", "字号"]
            found_keywords = [kw for kw in important_keywords if kw in truncated_part]
            
            if found_keywords:
                logger.error(f"⚠ 严重警告：截断部分包含重要关键词: {found_keywords}，信息将丢失！")
                print(f"[AI格式分析] ⚠ 严重警告：截断部分包含重要关键词: {found_keywords}，信息将丢失！")
                print("[AI格式分析] 建议：请检查文档，确保所有格式要求都在前8000字符内")
            
            # 截断时保留前8000字符，并添加警告标记
            document_text = document_text[:8000] + "\n\n[⚠️ 警告：文档内容过长，已截断。后续内容可能包含重要格式要求，请检查！]"
        
        # 构建提示词
        prompt_template_part1 = """你是文档格式分析专家。

## 完整指令系统（格式规范参考）：
"""
        prompt_template_part2 = """

## 格式要求文档内容：
"""
        prompt_template_part3 = """

## 任务：
分析上述格式要求文档，理解其中的格式规范，然后：
1. 生成清晰、准确的自然语言格式描述（用于展示给用户）
2. 根据完整指令系统的结构，生成JSON格式的子集指令系统（用于系统执行）

## 要求：
1. 自然语言描述应该包括：
   - 正文格式（字体、字号、行距、首行缩进等）
   - 标题格式（各级标题的字体、字号、对齐方式、前后间距等）
   - 页面设置（页边距、纸张大小等）
   - 特殊格式（目录、摘要、关键词、结论、参考文献等格式要求）
   - 章节编号格式（如：第一章、1.1等）
   - 分页规则（如：章节是否从新页开始）

2. JSON格式指令必须：
   - 严格按照完整指令系统的JSON结构
   - 只包含文档中实际提到的格式配置
   - 字段名必须与完整指令系统一致
   - 字段值必须在完整指令系统定义的允许范围内
   - 优先使用空行数字段（*_before_lines、*_after_lines）而不是磅数字段

3. **重要格式要求（必须严格遵守）**：
   - **目录标题文本**：如果文档中明确提到"目录"或"目 录"，必须使用 `"目　　录"`（两个全角空格，Unicode: \u3000\u3000），不要使用 `"目 录"`（半角空格）或 `"目录"`（无空格）
   - **结论标题文本**：如果文档中明确提到"结论"或"结 论"，必须使用 `"结　　论"`（两个全角空格，Unicode: \u3000\u3000），不要使用 `"结 论"`（半角空格）或 `"结论"`（无空格）
   - **一级标题字号**：必须准确识别文档中实际指定的一级标题字号。如果文档明确提到"小三号"、"15磅"或类似表述，必须使用 `font_size_pt: 15`，不要使用默认的 14

4. 如果文档中存在完整指令系统中没有定义的格式要求，请在自然语言描述中明确说明，并在JSON中使用extended_fields记录。

## 输出格式：
请按照以下格式输出：

【自然语言格式描述】
（这里写自然语言格式描述）

【JSON格式指令】
```json
{
  "version": "...",
  "format_rules": { ... },
  "application_rules": { ... }
}
```

请开始分析并生成结果。"""
        
        prompt = (
            prompt_template_part1 + universal_system_json +
            prompt_template_part2 + document_text +
            prompt_template_part3
        )
        
        # 输出提交给AI的完整提示词（用于调试，确保文档文本完整传递）
        print("=" * 100)
        print("[AI格式分析] 📤 提交给AI的完整提示词：")
        print("=" * 100)
        prompt_preview_length = 2000  # 显示前2000字符
        if len(prompt) > prompt_preview_length:
            print(prompt[:prompt_preview_length])
            print(f"\n... [提示词过长，已截断前{prompt_preview_length}字符，总长度: {len(prompt)}字符] ...")
        else:
            print(prompt)
        print("=" * 100)
        logger.info("=" * 100)
        logger.info("[AI格式分析] 📤 提交给AI的完整提示词：")
        logger.info("=" * 100)
        if len(prompt) > prompt_preview_length:
            logger.info(prompt[:prompt_preview_length] + f"\n... [提示词过长，已截断前{prompt_preview_length}字符，总长度: {len(prompt)}字符] ...")
        else:
            logger.info(prompt)
        logger.info("=" * 100)
        
        # 特别检查文档文本部分是否包含"空两行"等关键词
        if "空两行" in document_text or "空两行" in prompt:
            print("[AI格式分析] ✓ 检测到文档中包含'空两行'文字说明")
            logger.info("[AI格式分析] ✓ 检测到文档中包含'空两行'文字说明")
        if "空行" in document_text or "空行" in prompt:
            print("[AI格式分析] ✓ 检测到文档中包含'空行'相关文字说明")
            logger.info("[AI格式分析] ✓ 检测到文档中包含'空行'相关文字说明")
        
        # 获取AI提供商
        print(f"[AI格式分析] 正在获取AI模型配置...")
        logger.info(f"[AI格式分析] 正在获取AI模型配置...")
        llm_provider, model_config = await AiGenerationService._get_ai_provider(query_db, config_id)
        print(f"[AI格式分析] ✓ AI模型配置获取完成")
        print(f"  AI模型: {model_config.model_name if model_config and hasattr(model_config, 'model_name') else 'N/A'}")
        print(f"  配置ID: {config_id or '使用默认配置'}")
        logger.info(f"[AI格式分析] ✓ AI模型配置获取完成")
        logger.info(f"  AI模型: {model_config.model_name if model_config and hasattr(model_config, 'model_name') else 'N/A'}")
        logger.info(f"  配置ID: {config_id or '使用默认配置'}")
        
        messages = [
            {
                "role": "system",
                "content": "你是一位专业的文档格式分析专家，擅长分析格式要求文档并生成详细的格式化指令。"
            },
            {"role": "user", "content": prompt}
        ]
        
        print("=" * 100)
        print(f"[AI格式分析] 📤 正在发送请求给AI模型...")
        print(f"  开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"  提示词长度: {len(prompt)} 字符")
        print(f"  超时设置: 120秒")
        print("=" * 100)
        import sys
        sys.stdout.flush()
        logger.info("=" * 100)
        logger.info(f"[AI格式分析] 📤 正在发送请求给AI模型...")
        logger.info(f"  开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        logger.info(f"  提示词长度: {len(prompt)} 字符")
        logger.info(f"  超时设置: 120秒")
        logger.info("=" * 100)
        
        try:
            import asyncio
            ai_call_start = time.time()
            response = await asyncio.wait_for(
                llm_provider.chat(messages, temperature=0.3, max_tokens=4000),
                timeout=120.0
            )
            ai_call_elapsed = time.time() - ai_call_start
            print("=" * 100)
            print(f"[AI格式分析] ✓ AI模型响应接收完成")
            print(f"  完成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            print(f"  耗时: {ai_call_elapsed:.2f} 秒 ({ai_call_elapsed / 60:.2f} 分钟)")
            print(f"  响应类型: {type(response).__name__}")
            if hasattr(response, 'content'):
                response_length = len(str(response.content)) if response.content else 0
                print(f"  响应内容长度: {response_length} 字符")
            print("=" * 100)
            sys.stdout.flush()
            logger.info("=" * 100)
            logger.info(f"[AI格式分析] ✓ AI模型响应接收完成")
            logger.info(f"  完成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            logger.info(f"  耗时: {ai_call_elapsed:.2f} 秒 ({ai_call_elapsed / 60:.2f} 分钟)")
            logger.info(f"  响应类型: {type(response).__name__}")
            if hasattr(response, 'content'):
                response_length = len(str(response.content)) if response.content else 0
                logger.info(f"  响应内容长度: {response_length} 字符")
            logger.info("=" * 100)
            
            # 解析AI响应
            result = cls._parse_ai_format_response(response)
            natural_language = result.get('natural_language_description', '')
            json_str = result.get('json_instructions', '{}')
            
            # 解析JSON指令
            try:
                subset_instruction = json.loads(json_str) if isinstance(json_str, str) else json_str
            except json.JSONDecodeError:
                logger.warning("AI返回的JSON格式指令无效，尝试修复...")
                subset_instruction = cls._extract_json_from_text(json_str)
            
            return {
                'natural_language_description': natural_language,
                'subset_instruction': subset_instruction
            }
            
        except asyncio.TimeoutError:
            logger.error("AI调用超时（超过120秒）")
            raise ServiceException(message="AI分析格式文档超时，请稍后重试")
        except Exception as e:
            logger.error(f"AI分析失败: {str(e)}", exc_info=True)
            raise ServiceException(message=f"AI分析格式文档失败: {str(e)}")
    
    @classmethod
    async def _analyze_format_with_ai_simple(
        cls,
        query_db: AsyncSession,
        document_text: str,
        config_id: Optional[int] = None
    ) -> Dict[str, str]:
        """
        简化版AI分析方法（当没有完整指令系统时）
        
        :param query_db: 数据库会话
        :param document_text: 文档文本内容
        :param config_id: AI模型配置ID
        :return: 包含自然语言描述和JSON指令的字典
        """
        # 限制文档文本长度（避免超过AI上下文限制）
        # 重要：不能丢失任何文字信息！如果必须截断，要确保截断位置合理
        original_length = len(document_text)
        if original_length > 8000:
            logger.warning(f"文档文本过长（{original_length}字符），需要截断，但必须确保不丢失关键信息")
            print(f"[AI格式分析] ⚠ 警告：文档文本过长（{original_length}字符），需要截断")
            
            # 检查截断部分是否包含重要关键词
            truncated_part = document_text[8000:]
            important_keywords = ["空两行", "空行", "格式要求", "标题", "正文", "页边距", "字体", "字号"]
            found_keywords = [kw for kw in important_keywords if kw in truncated_part]
            
            if found_keywords:
                logger.error(f"⚠ 严重警告：截断部分包含重要关键词: {found_keywords}，信息将丢失！")
                print(f"[AI格式分析] ⚠ 严重警告：截断部分包含重要关键词: {found_keywords}，信息将丢失！")
                print("[AI格式分析] 建议：请检查文档，确保所有格式要求都在前8000字符内")
            
            # 截断时保留前8000字符，并添加警告标记
            document_text = document_text[:8000] + "\n\n[⚠️ 警告：文档内容过长，已截断。后续内容可能包含重要格式要求，请检查！]"
        
        # 使用字符串拼接而不是f-string，避免文档文本中的大括号导致解析错误
        prompt_template = """你是文档格式分析专家。

## 格式要求文档内容：
"""
        prompt_template_end = """

## 任务：
分析上述格式要求文档，理解其中的格式规范，然后：
1. 生成清晰、准确的自然语言格式描述
2. 生成JSON格式的格式化指令

请按照以下格式输出：

【自然语言格式描述】
（这里写自然语言格式描述）

【JSON格式指令】
```json
{
  "format_rules": { ... },
  "application_rules": { ... }
}
```"""
        prompt = prompt_template + document_text + prompt_template_end
        
        # 输出提交给AI的完整提示词（用于调试）
        print("=" * 100)
        print("[AI格式分析] 📤 提交给AI的完整提示词（简化版）：")
        print("=" * 100)
        prompt_preview_length = 2000
        if len(prompt) > prompt_preview_length:
            print(prompt[:prompt_preview_length])
            print(f"\n... [提示词过长，已截断前{prompt_preview_length}字符，总长度: {len(prompt)}字符] ...")
        else:
            print(prompt)
        print("=" * 100)
        logger.info("=" * 100)
        logger.info("[AI格式分析] 📤 提交给AI的完整提示词（简化版）：")
        logger.info("=" * 100)
        if len(prompt) > prompt_preview_length:
            logger.info(prompt[:prompt_preview_length] + f"\n... [提示词过长，已截断前{prompt_preview_length}字符，总长度: {len(prompt)}字符] ...")
        else:
            logger.info(prompt)
        logger.info("=" * 100)
        
        # 特别检查文档文本部分是否包含"空两行"等关键词
        if "空两行" in document_text or "空两行" in prompt:
            print("[AI格式分析] ✓ 检测到文档中包含'空两行'文字说明")
            logger.info("[AI格式分析] ✓ 检测到文档中包含'空两行'文字说明")
        if "空行" in document_text or "空行" in prompt:
            print("[AI格式分析] ✓ 检测到文档中包含'空行'相关文字说明")
            logger.info("[AI格式分析] ✓ 检测到文档中包含'空行'相关文字说明")
        
        llm_provider, model_config = await AiGenerationService._get_ai_provider(query_db, config_id)
        messages = [
            {
                "role": "system",
                "content": "你是一位专业的文档格式分析专家，擅长分析格式要求文档并生成详细的格式化指令。"
            },
            {"role": "user", "content": prompt}
        ]
        
        try:
            import asyncio
            response = await asyncio.wait_for(
                llm_provider.chat(messages, temperature=0.3, max_tokens=4000),
                timeout=120.0
            )
            
            result = cls._parse_ai_format_response(response)
            json_str = result.get('json_instructions', '{}')
            
            try:
                subset_instruction = json.loads(json_str) if isinstance(json_str, str) else json_str
            except json.JSONDecodeError:
                subset_instruction = cls._extract_json_from_text(json_str)
            
            # 验证和修正
            try:
                validated_config = cls._validate_and_fix_format_config(subset_instruction)
                subset_instruction = validated_config
            except Exception as e:
                logger.warning(f"格式指令验证失败: {str(e)}，使用原始指令")
            
            return {
                'natural_language_description': result.get('natural_language_description', ''),
                'json_instructions': json.dumps(subset_instruction, ensure_ascii=False, indent=2)
            }
            
        except Exception as e:
            logger.error(f"AI分析失败: {str(e)}", exc_info=True)
            raise ServiceException(message=f"AI分析格式文档失败: {str(e)}")
    
    @classmethod
    async def _generate_natural_language_format_requirement(
        cls,
        query_db: AsyncSession,
        document_content: Dict[str, Any],
        universal_instruction_system: Dict[str, Any],
        config_id: Optional[int] = None
    ) -> str:
        """
        AI第一步：生成自然语言的格式要求
        
        :param query_db: 数据库会话
        :param document_content: 文档内容
        :param universal_instruction_system: 完整指令系统
        :param config_id: AI模型配置ID
        :return: 自然语言格式要求描述
        """
        # 将文档内容转换为JSON字符串
        content_json = json.dumps(document_content, ensure_ascii=False, indent=2)
        if len(content_json) > 3000:
            content_dict = document_content.copy()
            if 'paragraphs' in content_dict:
                for para in content_dict['paragraphs']:
                    if 'text' in para and len(para['text']) > 100:
                        para['text'] = para['text'][:100] + '...'
            content_json = json.dumps(content_dict, ensure_ascii=False, indent=2)[:3000]
        
        # 将完整指令系统转换为JSON字符串（只包含结构，不包含所有值）
        universal_system_json = json.dumps(universal_instruction_system, ensure_ascii=False, indent=2)
        if len(universal_system_json) > 5000:
            # 如果太长，只保留关键结构
            simplified_system = {
                'version': universal_instruction_system.get('version'),
                'description': universal_instruction_system.get('description'),
                'format_rules': {
                    'default_font': universal_instruction_system.get('format_rules', {}).get('default_font', {}),
                    'headings': universal_instruction_system.get('format_rules', {}).get('headings', {}),
                    'paragraph': universal_instruction_system.get('format_rules', {}).get('paragraph', {})
                },
                'application_rules': universal_instruction_system.get('application_rules', {})
            }
            universal_system_json = json.dumps(simplified_system, ensure_ascii=False, indent=2)
        
        # 构建第一步提示词
        prompt = f"""你是文档格式分析专家。

## 完整指令系统（参考）：
{universal_system_json}

## Word文档格式信息：
{content_json}

## 任务：
分析Word文档格式，生成自然语言的格式要求描述。

描述应该包括：
- 正文格式（字体、字号、行距、首行缩进等）
- 标题格式（各级标题的字体、字号、对齐方式、前后间距等）
- 页面设置（页边距、纸张大小等）
- 特殊格式（目录、摘要、关键词、结论、参考文献等格式要求）
- 章节编号格式（如：第一章、1.1等）
- 分页规则（如：章节是否从新页开始）

## 重要提示：
如果文档中存在完整指令系统中没有定义的格式要求（例如：特殊的字体、特殊的页眉页脚格式、特殊的表格样式等），请在自然语言描述中明确说明这些特殊格式要求，并标注为"扩展格式要求"。这些扩展格式要求将在后续步骤中通过扩展字段机制处理。

请用清晰、准确的自然语言描述格式要求，确保描述完整、准确。"""
        
        # 获取AI提供商
        print(f"[AI第一步] 获取AI提供商...")
        logger.info(f"[AI第一步] 获取AI提供商...")
        llm_provider, model_config = await AiGenerationService._get_ai_provider(query_db, config_id)
        provider_name = type(llm_provider).__name__
        print(f"[AI第一步] ✓ AI提供商: {provider_name}")
        logger.info(f"[AI第一步] ✓ AI提供商: {provider_name}")
        
        # 调用AI生成自然语言格式要求
        prompt_len = len(prompt)
        print(f"[AI第一步] 提示词长度: {prompt_len} 字符")
        logger.info(f"[AI第一步] 提示词长度: {prompt_len} 字符")
        
        messages = [
            {
                "role": "system",
                "content": "你是一位专业的文档格式分析专家，擅长分析Word文档的格式要求并用自然语言描述。"
            },
            {"role": "user", "content": prompt}
        ]
        
        try:
            import asyncio
            print(f"[AI第一步] 开始调用AI模型（超时时间: 120秒）...")
            logger.info(f"[AI第一步] 开始调用AI模型（超时时间: 120秒）...")
            import time
            start_time = time.time()
            
            response = await asyncio.wait_for(
                llm_provider.chat(messages, temperature=0.3, max_tokens=2000),
                timeout=120.0
            )
            
            elapsed_time = time.time() - start_time
            response_len = len(response)
            print(f"[AI第一步] ✓ AI调用完成，耗时: {elapsed_time:.2f}秒，响应长度: {response_len} 字符")
            logger.info(f"[AI第一步] ✓ AI调用完成，耗时: {elapsed_time:.2f}秒，响应长度: {response_len} 字符")
            
            # 清理响应（移除可能的标记）
            natural_language = response.strip()
            # 移除可能的【格式要求描述】等标记
            import re
            natural_language = re.sub(r'【.*?】', '', natural_language).strip()
            
            final_len = len(natural_language)
            print(f"[AI第一步] ✓ 自然语言格式要求生成完成，最终长度: {final_len} 字符")
            print(f"[AI第一步] 前200字符预览: {natural_language[:200]}...")
            logger.info(f"AI第一步完成：生成自然语言格式要求，长度: {final_len} 字符")
            return natural_language
            
        except asyncio.TimeoutError:
            print(f"[AI第一步] ✗ AI调用超时（超过120秒）")
            logger.error(f"AI第一步失败：AI调用超时（超过120秒）")
            raise ServiceException(message="生成自然语言格式要求超时，请稍后重试")
        except Exception as e:
            print(f"[AI第一步] ✗ AI调用失败: {str(e)}")
            logger.error(f"AI第一步失败：生成自然语言格式要求失败: {str(e)}", exc_info=True)
            raise ServiceException(message=f"生成自然语言格式要求失败: {str(e)}")
    
    @classmethod
    async def _generate_subset_instruction_system(
        cls,
        query_db: AsyncSession,
        document_content: Dict[str, Any],
        natural_language_description: str,
        universal_instruction_system: Dict[str, Any],
        config_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        AI第二步：按指令系统格式，生成子集指令系统
        
        :param query_db: 数据库会话
        :param document_content: 文档内容
        :param natural_language_description: 自然语言格式要求
        :param universal_instruction_system: 完整指令系统
        :param config_id: AI模型配置ID
        :return: 子集指令系统（字典格式）
        """
        # 将文档内容转换为JSON字符串
        content_json = json.dumps(document_content, ensure_ascii=False, indent=2)
        if len(content_json) > 3000:
            content_dict = document_content.copy()
            if 'paragraphs' in content_dict:
                for para in content_dict['paragraphs']:
                    if 'text' in para and len(para['text']) > 100:
                        para['text'] = para['text'][:100] + '...'
            content_json = json.dumps(content_dict, ensure_ascii=False, indent=2)[:3000]
        
        # 将完整指令系统转换为JSON字符串
        universal_system_json = json.dumps(universal_instruction_system, ensure_ascii=False, indent=2)
        if len(universal_system_json) > 10000:
            # 如果太长，只保留关键结构
            simplified_system = {
                'version': universal_instruction_system.get('version'),
                'description': universal_instruction_system.get('description'),
                'instruction_type': universal_instruction_system.get('instruction_type'),
                'format_rules': universal_instruction_system.get('format_rules', {}),
                'application_rules': universal_instruction_system.get('application_rules', {})
            }
            universal_system_json = json.dumps(simplified_system, ensure_ascii=False, indent=2)
        
        # 构建第二步提示词
        # 注意：使用字符串拼接而不是f-string或format，避免JSON中的大括号被误解析
        prompt_template_part1 = """你是文档格式分析专家。

## 完整指令系统（格式规范）：
"""
        prompt_template_part2 = """

## Word文档格式信息：
"""
        prompt_template_part3 = """

## 自然语言格式要求：
"""
        prompt_template_part4 = """

## 任务：
根据自然语言格式要求，从完整指令系统中选择相应的配置项，生成这个学校的子集指令系统。

要求：
1. 子集指令系统必须符合完整指令系统的格式规范（JSON结构必须一致）
2. 只包含这个学校实际使用的格式配置（从Word文档中实际提取的值）
3. 所有值必须从Word文档中实际提取，不要使用默认值
4. 必须严格按照完整指令系统的JSON结构
5. 字段名必须与完整指令系统一致
6. 字段值必须在完整指令系统定义的允许范围内

## 扩展字段处理（重要）：
如果自然语言格式要求中提到了"扩展格式要求"或存在以下情况，必须使用扩展字段机制：
1. 格式要求在完整指令系统的format_rules或application_rules中找不到对应字段
2. 字段值不在完整指令系统定义的allowed_values范围内
3. 存在特殊的格式规则（如特殊的页眉页脚格式、特殊的表格样式等）

扩展字段结构：
```json
{
  "version": "...",
  "format_rules": { ... },
  "application_rules": { ... },
  "extended_fields": {
    "format_rules": {
      "字段名": {
        "name": "字段显示名称",
        "value": "实际值",
        "description": "字段说明",
        "source": "从Word文档的哪个部分提取的"
      }
    },
    "application_rules": {
      "字段名": {
        "name": "字段显示名称",
        "value": "实际值",
        "description": "字段说明",
        "source": "从Word文档的哪个部分提取的"
      }
    }
  }
}
```

扩展字段要求：
- 扩展字段名应使用下划线命名（snake_case），清晰描述字段用途
- 每个扩展字段必须包含name、value、description、source四个字段
- value应保持原始格式，不要转换
- description应说明该字段的用途和来源
- source应说明该值是从Word文档的哪个部分提取的

重要：
- 字体大小单位是"磅"（point），正常范围是8-30磅
- 如果文档中字体大小是数字（如 `12`），直接使用
- 如果提取的值不在正常范围内（如45.72磅），需要重新检查
- 所有格式必须从文档中实际提取，不使用默认值
- 优先使用完整指令系统中的字段，只有在找不到对应字段时才使用扩展字段

## 空行数字段使用规则（重要）：
对于标题和特殊章节的空行设置，优先使用空行数字段（*_before_lines、*_after_lines）：
1. **优先使用空行数字段**：如果Word文档中明确提到"空X行"、"前后各空两行"等描述，直接使用空行数字段（如：title_before_lines: 2, title_after_lines: 2）
2. **从磅数转换**：如果只有磅数信息（如：title_spacing_before_pt: 48），可以转换为空行数（通常24磅=1行），或同时提供两个字段
3. **字段优先级**：空行数字段优先于磅数字段。如果两者都有，优先使用空行数字段
4. **适用范围**：
   - 标题（h1/h2/h3）：使用 spacing_before_lines 和 spacing_after_lines
   - 特殊章节（目录、结论、参考文献、致谢等）：使用 title_before_lines 和 title_after_lines
   - 空行数范围：0-10行

请生成JSON格式的子集指令系统，严格按照完整指令系统的结构。如果存在扩展格式要求，请在extended_fields中记录。"""
        
        # 使用字符串拼接填充变量，避免f-string嵌套过深和format方法将JSON大括号误解析的问题
        prompt = (
            prompt_template_part1 + universal_system_json +
            prompt_template_part2 + content_json +
            prompt_template_part3 + natural_language_description +
            prompt_template_part4
        )
        
        # 获取AI提供商
        print(f"[AI第二步] 获取AI提供商...")
        logger.info(f"[AI第二步] 获取AI提供商...")
        llm_provider, model_config = await AiGenerationService._get_ai_provider(query_db, config_id)
        provider_name = type(llm_provider).__name__
        print(f"[AI第二步] ✓ AI提供商: {provider_name}")
        logger.info(f"[AI第二步] ✓ AI提供商: {provider_name}")
        
        # 调用AI生成子集指令系统
        prompt_len = len(prompt)
        print(f"[AI第二步] 提示词长度: {prompt_len} 字符")
        print(f"[AI第二步] 自然语言描述长度: {len(natural_language_description)} 字符")
        logger.info(f"[AI第二步] 提示词长度: {prompt_len} 字符")
        logger.info(f"[AI第二步] 自然语言描述长度: {len(natural_language_description)} 字符")
        
        messages = [
            {
                "role": "system",
                "content": "你是一位专业的文档格式分析专家，擅长根据格式要求生成符合规范的JSON格式指令。"
            },
            {"role": "user", "content": prompt}
        ]
        
        try:
            import asyncio
            print(f"[AI第二步] 开始调用AI模型（超时时间: 120秒）...")
            logger.info(f"[AI第二步] 开始调用AI模型（超时时间: 120秒）...")
            import time
            start_time = time.time()
            
            response = await asyncio.wait_for(
                llm_provider.chat(messages, temperature=0.3, max_tokens=4000),
                timeout=120.0
            )
            
            elapsed_time = time.time() - start_time
            response_len = len(response)
            print(f"[AI第二步] ✓ AI调用完成，耗时: {elapsed_time:.2f}秒，响应长度: {response_len} 字符")
            logger.info(f"[AI第二步] ✓ AI调用完成，耗时: {elapsed_time:.2f}秒，响应长度: {response_len} 字符")
            
            # 解析AI响应，提取JSON指令
            print(f"[AI第二步] 开始解析AI响应，提取JSON指令...")
            logger.info(f"[AI第二步] 开始解析AI响应，提取JSON指令...")
            # 尝试提取JSON部分
            import re
            json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response, re.DOTALL)
            if json_match:
                json_str = json_match.group(1).strip()
                print(f"[AI第二步] ✓ 从代码块中提取到JSON")
                logger.info(f"[AI第二步] ✓ 从代码块中提取到JSON")
            else:
                # 尝试从响应中提取JSON对象
                print(f"[AI第二步] 未找到代码块，尝试从文本中提取JSON...")
                logger.info(f"[AI第二步] 未找到代码块，尝试从文本中提取JSON...")
                json_str = cls._extract_json_from_text(response)
                print(f"[AI第二步] ✓ 从文本中提取到JSON")
                logger.info(f"[AI第二步] ✓ 从文本中提取到JSON")
            
            # 解析JSON
            if isinstance(json_str, str):
                try:
                    subset_instruction = json.loads(json_str)
                    print(f"[AI第二步] ✓ JSON解析成功")
                    logger.info(f"[AI第二步] ✓ JSON解析成功")
                except json.JSONDecodeError as e:
                    print(f"[AI第二步] ⚠ JSON解析失败: {str(e)}，尝试使用_extract_json_from_text方法")
                    logger.warning(f"[AI第二步] JSON解析失败: {str(e)}，尝试使用_extract_json_from_text方法")
                    # 如果解析失败，尝试使用_extract_json_from_text方法
                    subset_instruction = cls._extract_json_from_text(json_str)
                    print(f"[AI第二步] ✓ 使用_extract_json_from_text方法成功")
                    logger.info(f"[AI第二步] ✓ 使用_extract_json_from_text方法成功")
            else:
                subset_instruction = json_str
                print(f"[AI第二步] ✓ 响应已经是字典格式")
                logger.info(f"[AI第二步] ✓ 响应已经是字典格式")
            
            instruction_size = len(json.dumps(subset_instruction, ensure_ascii=False))
            print(f"[AI第二步] ✓ 子集指令系统生成完成，大小: {instruction_size} 字符")
            logger.info(f"AI第二步完成：生成子集指令系统，大小: {instruction_size} 字符")
            return subset_instruction
            
        except asyncio.TimeoutError:
            print(f"[AI第二步] ✗ AI调用超时（超过120秒）")
            logger.error(f"AI第二步失败：AI调用超时（超过120秒）")
            raise ServiceException(message="生成子集指令系统超时，请稍后重试")
        except Exception as e:
            print(f"[AI第二步] ✗ AI调用失败: {str(e)}")
            logger.error(f"AI第二步失败：生成子集指令系统失败: {str(e)}", exc_info=True)
            raise ServiceException(message=f"生成子集指令系统失败: {str(e)}")
    
    @classmethod
    def _validate_format_specification(
        cls,
        subset_instruction: Dict[str, Any],
        universal_instruction_system: Dict[str, Any]
    ) -> list:
        """
        格式规范校验：检查子集指令系统是否符合完整指令系统的格式规范
        
        :param subset_instruction: 子集指令系统
        :param universal_instruction_system: 完整指令系统
        :return: 错误列表
        """
        errors = []
        
        if not universal_instruction_system:
            logger.warning("完整指令系统为空，跳过格式规范校验")
            return errors
        
        try:
            # 1. 检查JSON结构是否符合规范
            # 检查顶层字段
            required_top_level_fields = ['version', 'format_rules', 'application_rules']
            for field in required_top_level_fields:
                if field not in subset_instruction:
                    errors.append(f"缺少必填字段: {field}")
            
            # 2. 检查format_rules中的字段
            if 'format_rules' in subset_instruction and 'format_rules' in universal_instruction_system:
                subset_format_rules = subset_instruction['format_rules']
                universal_format_rules = universal_instruction_system['format_rules']
                
                # 检查每个字段是否在完整指令系统中定义
                for key, value in subset_format_rules.items():
                    if key not in universal_format_rules:
                        # 如果字段不在完整指令系统中，检查是否在扩展字段中
                        if 'extended_fields' in subset_instruction:
                            extended_format_rules = subset_instruction['extended_fields'].get('format_rules', {})
                            if key in extended_format_rules:
                                # 字段在扩展字段中，这是允许的，记录信息但不作为错误
                                logger.info(f"format_rules中的字段 '{key}' 使用扩展字段机制（符合规范）")
                                continue
                        # 如果既不在完整指令系统中，也不在扩展字段中，记录警告（不是错误）
                        logger.warning(f"format_rules中包含未定义的字段: {key}（建议使用扩展字段机制）")
                        # 不将其作为错误，因为可能是合理的扩展字段，只是没有正确放在extended_fields中
                        continue
                    
                    # 检查字段值是否符合规范
                    field_def = universal_format_rules[key]
                    if isinstance(field_def, dict) and 'type' in field_def:
                        # 如果字段定义包含类型信息，进行类型检查
                        field_type = field_def.get('type')
                        if field_type == 'number' and not isinstance(value, (int, float)):
                            errors.append(f"format_rules.{key} 应该是数字类型，实际为: {type(value).__name__}")
                        elif field_type == 'string' and not isinstance(value, str):
                            errors.append(f"format_rules.{key} 应该是字符串类型，实际为: {type(value).__name__}")
                        elif field_type == 'boolean' and not isinstance(value, bool):
                            errors.append(f"format_rules.{key} 应该是布尔类型，实际为: {type(value).__name__}")
                        
                        # 检查允许的值范围
                        if 'allowed_values' in field_def:
                            if isinstance(value, str) and value not in field_def['allowed_values']:
                                errors.append(f"format_rules.{key} 的值 '{value}' 不在允许的范围内: {field_def['allowed_values']}")
                        
                        # 检查数值范围
                        if 'range' in field_def and isinstance(value, (int, float)):
                            min_val, max_val = field_def['range']
                            if value < min_val or value > max_val:
                                errors.append(f"format_rules.{key} 的值 {value} 不在允许的范围内 [{min_val}, {max_val}]")
            
            # 3. 检查application_rules中的字段
            if 'application_rules' in subset_instruction and 'application_rules' in universal_instruction_system:
                subset_app_rules = subset_instruction['application_rules']
                universal_app_rules = universal_instruction_system['application_rules']
                
                # 检查关键字段
                key_fields = ['chapter_numbering_format', 'special_section_format_rules', 'document_structure']
                for field in key_fields:
                    if field in subset_app_rules:
                        if field not in universal_app_rules:
                            # 如果字段不在完整指令系统中，检查是否在扩展字段中
                            if 'extended_fields' in subset_instruction:
                                extended_app_rules = subset_instruction['extended_fields'].get('application_rules', {})
                                if field in extended_app_rules:
                                    # 字段在扩展字段中，这是允许的，记录信息但不作为错误
                                    logger.info(f"application_rules中的字段 '{field}' 使用扩展字段机制（符合规范）")
                                    continue
                            # 如果既不在完整指令系统中，也不在扩展字段中，记录警告
                            logger.warning(f"application_rules中包含未定义的字段: {field}（建议使用扩展字段机制）")
                            # 不将其作为错误，因为可能是合理的扩展字段
            
            if errors:
                logger.warning(f"格式规范校验发现 {len(errors)} 个问题: {errors}")
            else:
                logger.info("格式规范校验通过")
                
        except Exception as e:
            logger.error(f"格式规范校验过程出错: {str(e)}", exc_info=True)
            errors.append(f"格式规范校验过程出错: {str(e)}")
        
        return errors
    
    @classmethod
    def _validate_consistency(
        cls,
        natural_language: str,
        subset_instruction: Dict[str, Any]
    ) -> list:
        """
        一致性校验：检查自然语言描述与子集指令系统是否一致
        
        :param natural_language: 自然语言格式要求
        :param subset_instruction: 子集指令系统
        :return: 错误列表
        """
        errors = []
        
        try:
            import re
            
            # 1. 检查自然语言中提到的关键格式要求是否在子集指令系统中
            # 提取关键格式信息
            key_patterns = {
                'font_size': r'(\d+(?:\.\d+)?)\s*磅|字号[：:]\s*(\d+(?:\.\d+)?)',
                'font_name': r'字体[：:]\s*([^，,。\n]+)|([^，,。\n]+)体',
                'line_spacing': r'行距[：:]\s*(\d+(?:\.\d+)?)|(\d+(?:\.\d+)?)\s*倍行距',
                'alignment': r'对齐[：:]\s*([^，,。\n]+)|(居中|左对齐|右对齐)'
            }
            
            # 检查字体大小一致性
            font_size_matches = re.findall(key_patterns['font_size'], natural_language)
            if font_size_matches:
                for match in font_size_matches:
                    size_str = match[0] if match[0] else match[1]
                    if size_str:
                        try:
                            size_value = float(size_str)
                            # 检查子集指令系统中是否有对应的字体大小配置
                            found = False
                            if 'format_rules' in subset_instruction:
                                format_rules = subset_instruction['format_rules']
                                # 检查default_font
                                if 'default_font' in format_rules:
                                    default_font = format_rules['default_font']
                                    if 'size_pt' in default_font and abs(default_font['size_pt'] - size_value) < 0.1:
                                        found = True
                                # 检查headings
                                if not found and 'headings' in format_rules:
                                    for level in ['h1', 'h2', 'h3']:
                                        if level in format_rules['headings']:
                                            heading = format_rules['headings'][level]
                                            if 'font_size_pt' in heading and abs(heading['font_size_pt'] - size_value) < 0.1:
                                                found = True
                                                break
                            
                            if not found and size_value >= 8 and size_value <= 30:
                                # 如果字体大小在合理范围内但未在指令中找到，记录警告（不是错误）
                                logger.debug(f"自然语言中提到字体大小 {size_value} 磅，但在子集指令系统中未找到对应配置")
                        except ValueError:
                            pass
            
            # 2. 检查子集指令系统中的关键配置是否在自然语言中有描述
            # 检查字体大小
            if 'format_rules' in subset_instruction:
                format_rules = subset_instruction['format_rules']
                
                # 检查default_font
                if 'default_font' in format_rules:
                    default_font = format_rules['default_font']
                    if 'size_pt' in default_font:
                        size_value = default_font['size_pt']
                        if not re.search(rf'{size_value}\s*磅|字号[：:]\s*{size_value}', natural_language):
                            # 如果字体大小不在自然语言中，记录警告（不是错误，因为自然语言可能不完整）
                            logger.debug(f"子集指令系统中的默认字体大小 {size_value} 磅未在自然语言中明确提及")
            
            # 3. 检查扩展字段是否在自然语言中有描述
            if 'extended_fields' in subset_instruction:
                extended_fields = subset_instruction['extended_fields']
                
                # 检查format_rules扩展字段
                if 'format_rules' in extended_fields:
                    for field_name, field_info in extended_fields['format_rules'].items():
                        if isinstance(field_info, dict) and 'name' in field_info:
                            field_display_name = field_info.get('name', field_name)
                            # 检查自然语言中是否提到这个扩展字段
                            if field_display_name not in natural_language and field_name not in natural_language:
                                logger.warning(f"扩展字段 format_rules.{field_name} 未在自然语言格式要求中明确描述，建议补充说明")
                
                # 检查application_rules扩展字段
                if 'application_rules' in extended_fields:
                    for field_name, field_info in extended_fields['application_rules'].items():
                        if isinstance(field_info, dict) and 'name' in field_info:
                            field_display_name = field_info.get('name', field_name)
                            # 检查自然语言中是否提到这个扩展字段
                            if field_display_name not in natural_language and field_name not in natural_language:
                                logger.warning(f"扩展字段 application_rules.{field_name} 未在自然语言格式要求中明确描述，建议补充说明")
            
            # 4. 检查关键格式是否一致（更严格的检查）
            # 这里可以添加更多的一致性检查逻辑
            
            logger.info("一致性校验完成")
            
        except Exception as e:
            logger.error(f"一致性校验过程出错: {str(e)}", exc_info=True)
            errors.append(f"一致性校验过程出错: {str(e)}")
        
        return errors
    
    @classmethod
    def _validate_instruction_system(
        cls,
        natural_language: str,
        subset_instruction: Dict[str, Any],
        universal_instruction_system: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        整合三种校验方式：格式规范校验、一致性校验、数据质量校验
        
        :param natural_language: 自然语言格式要求
        :param subset_instruction: 子集指令系统
        :param universal_instruction_system: 完整指令系统
        :return: 校验结果
        """
        print("[指令系统校验] 开始执行三种校验...")
        logger.info("[指令系统校验] 开始执行三种校验...")
        all_errors = []
        
        # 1. 格式规范校验
        print("[指令系统校验] 步骤1/3: 格式规范校验...")
        logger.info("[指令系统校验] 步骤1/3: 格式规范校验...")
        format_errors = cls._validate_format_specification(subset_instruction, universal_instruction_system)
        all_errors.extend(format_errors)
        if format_errors:
            print(f"[指令系统校验] ⚠ 步骤1/3: 格式规范校验发现 {len(format_errors)} 个错误")
            logger.warning(f"[指令系统校验] ⚠ 步骤1/3: 格式规范校验发现 {len(format_errors)} 个错误")
            for i, error in enumerate(format_errors[:5], 1):  # 只显示前5个错误
                print(f"  错误{i}: {error}")
                logger.warning(f"  错误{i}: {error}")
        else:
            print(f"[指令系统校验] ✓ 步骤1/3: 格式规范校验通过")
            logger.info(f"[指令系统校验] ✓ 步骤1/3: 格式规范校验通过")
        
        # 2. 一致性校验
        print("[指令系统校验] 步骤2/3: 一致性校验...")
        logger.info("[指令系统校验] 步骤2/3: 一致性校验...")
        consistency_errors = cls._validate_consistency(natural_language, subset_instruction)
        all_errors.extend(consistency_errors)
        if consistency_errors:
            print(f"[指令系统校验] ⚠ 步骤2/3: 一致性校验发现 {len(consistency_errors)} 个错误")
            logger.warning(f"[指令系统校验] ⚠ 步骤2/3: 一致性校验发现 {len(consistency_errors)} 个错误")
            for i, error in enumerate(consistency_errors[:5], 1):  # 只显示前5个错误
                print(f"  错误{i}: {error}")
                logger.warning(f"  错误{i}: {error}")
        else:
            print(f"[指令系统校验] ✓ 步骤2/3: 一致性校验通过")
            logger.info(f"[指令系统校验] ✓ 步骤2/3: 一致性校验通过")
        
        # 3. 数据质量校验（复用现有的方法）
        print("[指令系统校验] 步骤3/3: 数据质量校验...")
        logger.info("[指令系统校验] 步骤3/3: 数据质量校验...")
        try:
            # 数据质量校验会自动修正异常值，所以这里只记录警告
            validated_config = cls._validate_and_fix_format_config(subset_instruction.copy())
            if validated_config != subset_instruction:
                print(f"[指令系统校验] ⚠ 步骤3/3: 数据质量校验发现并修正了异常值")
                logger.info("数据质量校验发现并修正了异常值")
            else:
                print(f"[指令系统校验] ✓ 步骤3/3: 数据质量校验通过")
                logger.info(f"[指令系统校验] ✓ 步骤3/3: 数据质量校验通过")
        except Exception as e:
            print(f"[指令系统校验] ✗ 步骤3/3: 数据质量校验过程出错: {str(e)}")
            logger.warning(f"数据质量校验过程出错: {str(e)}")
            all_errors.append(f"数据质量校验过程出错: {str(e)}")
        
        total_errors = len(all_errors)
        if total_errors == 0:
            print(f"[指令系统校验] ✓ 所有校验通过")
            logger.info(f"[指令系统校验] ✓ 所有校验通过")
        else:
            print(f"[指令系统校验] ⚠ 共发现 {total_errors} 个错误")
            logger.warning(f"[指令系统校验] ⚠ 共发现 {total_errors} 个错误")
        
        return {
            'valid': len(all_errors) == 0,
            'errors': all_errors
        }
    
    @classmethod
    def _create_formatted_document(
        cls,
        chapters: list,
        format_config: Dict[str, Any],
        thesis_id: int,
        thesis = None,
        layout_rules: Dict[str, Any] = None
    ) -> str:
        """
        创建格式化的Word文档
        
        :param chapters: 章节列表（已完成状态）
        :param format_config: 格式配置
        :param thesis_id: 论文ID
        :param thesis: 论文对象（可选，用于添加标题等信息）
        :return: 输出文件路径
        """
        from utils.log_util import logger
        
        if not DOCX_AVAILABLE:
            import sys
            python_path = sys.executable
            error_msg = (
                f'python-docx 未安装，无法创建Word文档。\n'
                f'当前Python路径: {python_path}\n'
                f'请执行以下命令安装: pip install python-docx\n'
                f'如果使用虚拟环境，请确保已激活虚拟环境。'
            )
            logger.error(error_msg)
            raise ServiceException(message=error_msg)
        
        try:
            logger.info(f"[格式化开始] 论文ID: {thesis_id}, 章节数量: {len(chapters)}")
            
            # 创建新文档
            logger.info(f"[步骤1/6] 创建新Word文档")
            doc = Document()
            
            # 应用页面设置
            logger.info(f"[步骤2/6] 应用页面设置")
            if 'page' in format_config:
                page_config = format_config['page']
                if 'margins' in page_config:
                    margins = page_config['margins']
                    section = doc.sections[0]
                    section.top_margin = Inches(float(margins.get('top', 72)) / 72)
                    section.bottom_margin = Inches(float(margins.get('bottom', 72)) / 72)
                    section.left_margin = Inches(float(margins.get('left', 90)) / 72)
                    section.right_margin = Inches(float(margins.get('right', 90)) / 72)
                    logger.info(f"  页边距设置: 上={margins.get('top', 72)}磅, 下={margins.get('bottom', 72)}磅, 左={margins.get('left', 90)}磅, 右={margins.get('right', 90)}磅")
            
            # 获取格式配置
            logger.info(f"[步骤3/6] 解析格式配置")
            font_config = format_config.get('font', {})
            para_config = format_config.get('paragraph', {})
            headings_config = format_config.get('headings', {})
            
            # 设置默认字体
            default_font_name = font_config.get('name', '宋体')
            default_font_size = Pt(font_config.get('size', 12))
            logger.info(f"  默认字体: {default_font_name}, 字号: {font_config.get('size', 12)}磅")
            logger.info(f"  段落格式: 行距={para_config.get('line_spacing', 1.5)}, 首行缩进={para_config.get('first_line_indent', 24)}磅")
            logger.info(f"  标题格式: h1={headings_config.get('h1', {}).get('font_size', 'N/A')}磅, h2={headings_config.get('h2', {}).get('font_size', 'N/A')}磅, h3={headings_config.get('h3', {}).get('font_size', 'N/A')}磅")
            
            # 获取布局规则
            if layout_rules is None:
                layout_rules = {}
            
            title_spacing = layout_rules.get('title_spacing', {})
            heading_spacing = layout_rules.get('heading_spacing', {})
            chapter_spacing = layout_rules.get('chapter_spacing', {})
            section_spacing = layout_rules.get('section_spacing', {})
            paragraph_spacing = layout_rules.get('paragraph_spacing', {})
            logger.info(f"  布局规则: 标题后空行={title_spacing.get('after_title', 1)}, 章节间距={chapter_spacing.get('between_chapters', 'page_break')}")
            
            # 添加论文标题（如果有）
            if thesis and thesis.title:
                logger.info(f"[步骤4/6] 添加论文标题: {thesis.title}")
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(thesis.title)
                title_run.font.name = default_font_name
                title_run.font.size = Pt(18)  # 标题字体稍大
                title_run.font.bold = True
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 应用标题后的空行规则
                after_title_lines = title_spacing.get('after_title', 1)
                for _ in range(after_title_lines):
                    doc.add_paragraph()
                logger.info(f"  标题后添加 {after_title_lines} 个空行")
            
            # 检查是否需要生成目录
            has_toc_chapter = any(
                '目录' in (getattr(ch, 'title', '') or '') or 
                'table of contents' in (getattr(ch, 'title', '') or '').lower() or
                'toc' in (getattr(ch, 'title', '') or '').lower()
                for ch in chapters
            )
            
            # 从格式配置中读取是否自动生成目录
            application_rules = format_config.get('application_rules', {})
            should_generate_toc = application_rules.get('auto_generate_toc', False)
            
            # 如果需要生成目录且没有目录章节，则自动生成
            if should_generate_toc and not has_toc_chapter:
                logger.info(f"[步骤5/6] 检测到需要自动生成目录")
                try:
                    toc_chapter = cls._generate_table_of_contents(chapters, format_config, layout_rules)
                    if toc_chapter and hasattr(toc_chapter, 'toc_entries') and len(toc_chapter.toc_entries) > 0:
                        # 将目录插入到第一个位置（在摘要之前）
                        chapters.insert(0, toc_chapter)
                        logger.info(f"  已自动生成目录章节，共 {len(toc_chapter.toc_entries)} 个条目")
                except Exception as e:
                    logger.warning(f"  自动生成目录失败: {str(e)}，将跳过目录生成")
            
            # 按章节顺序排序（确保与大纲顺序一致）
            chapters = sorted(chapters, key=lambda x: getattr(x, 'order_num', 0) if hasattr(x, 'order_num') else 0)
            logger.info(f"[步骤5/6] 开始处理章节内容，共 {len(chapters)} 个章节")
            
            # 遍历章节，添加内容
            for idx, chapter in enumerate(chapters):
                if not chapter.content:
                    logger.warning(f"  章节 {idx+1}: {chapter.title} - 内容为空，跳过")
                    continue
                
                logger.info(f"  处理章节 {idx+1}/{len(chapters)}: {chapter.title} (级别: {getattr(chapter, 'level', 1)})")
                
                # 章节之间的间距处理（除了第一个章节）
                if idx > 0:
                    between_chapters = chapter_spacing.get('between_chapters', 'page_break')
                    if between_chapters == 'page_break':
                        doc.add_page_break()
                        logger.debug(f"    章节之间添加分页符")
                    elif isinstance(between_chapters, int) and between_chapters > 0:
                        for _ in range(between_chapters):
                            doc.add_paragraph()
                        logger.debug(f"    章节之间添加 {between_chapters} 个空行")
                
                # 章节开始前的空行
                before_chapter = chapter_spacing.get('before_chapter', 0)
                for _ in range(before_chapter):
                    doc.add_paragraph()
                if before_chapter > 0:
                    logger.debug(f"    章节前添加 {before_chapter} 个空行")
                
                # 添加章节标题
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(chapter.title)
                title_run.bold = True
                
                # 应用标题格式（根据章节级别对应格式指令）
                # 章节级别（chapter.level）对应格式指令中的 headings.h{level}
                # 例如：level=1 → headings.h1, level=2 → headings.h2, level=3 → headings.h3
                chapter_level = chapter.level if hasattr(chapter, 'level') else 1
                level_key = f'h{chapter_level}'
                
                if headings_config and level_key in headings_config:
                    heading_style = headings_config[level_key]
                    title_run.font.size = Pt(heading_style.get('font_size', 16))
                    title_run.font.name = heading_style.get('font_name', default_font_name)
                    title_run.font.bold = heading_style.get('bold', True)
                    logger.info(f"    应用标题格式: {heading_style.get('font_name', default_font_name)} {heading_style.get('font_size', 16)}磅, 加粗={heading_style.get('bold', True)}")
                else:
                    # 如果没有对应级别的标题格式，使用默认值
                    title_run.font.size = Pt(16)
                    title_run.font.name = default_font_name
                    title_run.font.bold = True
                    logger.warning(f"    未找到级别 {chapter_level} 的标题格式，使用默认格式")
                
                # 设置标题对齐（从格式指令中获取）
                title_alignment = headings_config.get(level_key, {}).get('alignment', 'center') if headings_config else 'center'
                if title_alignment == 'center':
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif title_alignment == 'right':
                    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 默认居中
                
                # 设置标题段间距（从格式指令中获取）
                if headings_config and level_key in headings_config:
                    heading_style = headings_config[level_key]
                    if 'spacing_before' in heading_style:
                        title_para.paragraph_format.space_before = Pt(heading_style['spacing_before'])
                    if 'spacing_after' in heading_style:
                        title_para.paragraph_format.space_after = Pt(heading_style['spacing_after'])
                
                # 应用标题后的空行规则（从layout_rules中获取）
                heading_spacing_config = heading_spacing.get(level_key, {})
                after_heading_lines = heading_spacing_config.get('after', 1)  # 默认空1行
                for _ in range(after_heading_lines):
                    doc.add_paragraph()
                if after_heading_lines > 0:
                    logger.debug(f"    标题后添加 {after_heading_lines} 个空行")
                
                # 获取特殊格式配置
                special_formats = format_config.get('special_formats', {})
                
                # 识别章节类型（用于应用特殊格式和布局规则）
                chapter_title_lower = chapter.title.lower() if chapter.title else ''
                is_special_section = False
                special_type = None
                
                # 检查是否是特殊章节（目录、摘要、关键词、结论）
                if '目录' in chapter.title or 'table of contents' in chapter_title_lower or 'toc' in chapter_title_lower:
                    is_special_section = True
                    special_type = 'table_of_contents'
                    logger.info(f"    识别为特殊章节: 目录")
                    
                    # 如果是自动生成的目录章节，生成目录内容
                    if hasattr(chapter, 'is_toc') and chapter.is_toc and hasattr(chapter, 'toc_entries'):
                        logger.info(f"    生成目录内容，共 {len(chapter.toc_entries)} 个条目")
                        # 生成目录内容（Markdown格式）
                        toc_content_lines = []
                        for entry in chapter.toc_entries:
                            level = entry.get('level', 1)
                            title = entry.get('title', '')
                            page_num = entry.get('page_number', '')
                            
                            # 根据级别添加缩进
                            indent = '  ' * (level - 1)
                            
                            # 生成目录行（格式：标题 ... 页码）
                            if page_num:
                                toc_line = f"{indent}{title} ... {page_num}"
                            else:
                                toc_line = f"{indent}{title}"
                            
                            toc_content_lines.append(toc_line)
                        
                        # 将生成的目录内容赋值给章节内容
                        chapter.content = '\n'.join(toc_content_lines)
                        logger.info(f"    目录内容已生成，共 {len(toc_content_lines)} 行")
                elif '摘要' in chapter.title or 'abstract' in chapter_title_lower:
                    is_special_section = True
                    special_type = 'abstract'
                    logger.info(f"    识别为特殊章节: 摘要")
                elif '关键词' in chapter.title or 'keywords' in chapter_title_lower:
                    is_special_section = True
                    special_type = 'keywords'
                    logger.info(f"    识别为特殊章节: 关键词")
                elif '结论' in chapter.title or 'conclusion' in chapter_title_lower:
                    is_special_section = True
                    special_type = 'conclusion'
                    logger.info(f"    识别为特殊章节: 结论")
                
                # 应用特殊章节前的布局规则（在标题之前）
                if is_special_section and special_type in section_spacing:
                    section_spacing_config = section_spacing[special_type]
                    before_section = section_spacing_config.get('before', 0)
                    # 注意：这里需要在标题之前添加空行，但标题已经添加了，所以这里先记录，稍后处理
                    # 实际上，特殊章节的before应该在章节开始前处理，已在chapter_spacing中处理
                
                # 处理章节内容（支持Markdown格式）
                content_lines = chapter.content.split('\n')
                total_lines = len(content_lines)
                logger.info(f"    开始处理章节内容，共 {total_lines} 行")
                current_heading_level = None
                is_first_line = True  # 用于识别特殊章节的第一行（可能是标题）
                
                # 如果是目录章节，记录需要在第一行（目录标题）前后添加的空行数
                toc_title_before_lines = 0
                toc_title_after_lines = 0
                if is_special_section and special_type == 'table_of_contents':
                    toc_spacing_config = section_spacing.get('table_of_contents', {})
                    toc_title_before_lines = toc_spacing_config.get('title_before', 0)
                    toc_title_after_lines = toc_spacing_config.get('title_after', 0)
                    if toc_title_before_lines > 0 or toc_title_after_lines > 0:
                        logger.info(f"    目录标题空行规则: 前{toc_title_before_lines}行, 后{toc_title_after_lines}行")
                
                processed_lines = 0
                markdown_headings_count = 0
                paragraphs_count = 0
                
                for line_num, line in enumerate(content_lines, 1):
                    line = line.strip()
                    if not line:
                        # 空行
                        doc.add_paragraph()
                        continue
                    
                    processed_lines += 1
                    
                    # 如果是目录章节的第一行，且需要添加标题前的空行
                    if is_special_section and special_type == 'table_of_contents' and is_first_line and toc_title_before_lines > 0:
                        for _ in range(toc_title_before_lines):
                            doc.add_paragraph()
                        logger.debug(f"      第{line_num}行前: 添加目录标题前空行 {toc_title_before_lines} 行")
                    
                    # 检测Markdown标题（## 或 ###）
                    if line.startswith('##'):
                            # 二级标题
                        markdown_headings_count += 1
                        heading_text = line.lstrip('#').strip()
                        logger.debug(f"      第{line_num}行: Markdown二级标题 - {heading_text}")
                        heading_para = doc.add_paragraph()
                        heading_run = heading_para.add_run(heading_text)
                        
                        # 应用标题格式（使用h2配置）
                        if headings_config and 'h2' in headings_config:
                            h2_style = headings_config['h2']
                            heading_run.font.size = Pt(h2_style.get('font_size', 14))
                            heading_run.font.name = h2_style.get('font_name', '黑体')
                            heading_run.font.bold = h2_style.get('bold', True)
                        else:
                            heading_run.font.size = Pt(14)
                            heading_run.font.name = '黑体'
                            heading_run.font.bold = True
                        
                        heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        if headings_config and 'h2' in headings_config:
                            h2_style = headings_config['h2']
                            if 'spacing_before' in h2_style:
                                heading_para.paragraph_format.space_before = Pt(h2_style['spacing_before'])
                            if 'spacing_after' in h2_style:
                                heading_para.paragraph_format.space_after = Pt(h2_style['spacing_after'])
                        else:
                            heading_para.paragraph_format.space_before = Pt(12)
                            heading_para.paragraph_format.space_after = Pt(6)
                        
                        # 应用二级标题后的空行规则（从layout_rules中获取）
                        h2_spacing_config = heading_spacing.get('h2', {})
                        after_h2_lines = h2_spacing_config.get('after', 0)  # 默认不空行
                        for _ in range(after_h2_lines):
                            doc.add_paragraph()
                        continue
                    elif line.startswith('###'):
                        # 三级标题
                        markdown_headings_count += 1
                        heading_text = line.lstrip('#').strip()
                        logger.debug(f"      第{line_num}行: Markdown三级标题 - {heading_text}")
                        heading_para = doc.add_paragraph()
                        heading_run = heading_para.add_run(heading_text)
                        
                        # 应用标题格式（使用h3配置）
                        if headings_config and 'h3' in headings_config:
                            h3_style = headings_config['h3']
                            heading_run.font.size = Pt(h3_style.get('font_size', 12))
                            heading_run.font.name = h3_style.get('font_name', '黑体')
                            heading_run.font.bold = h3_style.get('bold', True)
                        else:
                            heading_run.font.size = Pt(12)
                            heading_run.font.name = '黑体'
                            heading_run.font.bold = True
                        
                        heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        if headings_config and 'h3' in headings_config:
                            h3_style = headings_config['h3']
                            if 'spacing_before' in h3_style:
                                heading_para.paragraph_format.space_before = Pt(h3_style['spacing_before'])
                            if 'spacing_after' in h3_style:
                                heading_para.paragraph_format.space_after = Pt(h3_style['spacing_after'])
                        else:
                            heading_para.paragraph_format.space_before = Pt(10)
                            heading_para.paragraph_format.space_after = Pt(4)
                        
                        # 应用三级标题后的空行规则（从layout_rules中获取）
                        h3_spacing_config = heading_spacing.get('h3', {})
                        after_h3_lines = h3_spacing_config.get('after', 0)  # 默认不空行
                        for _ in range(after_h3_lines):
                            doc.add_paragraph()
                        continue
                    
                    # 处理普通段落（支持Markdown加粗）
                    paragraphs_count += 1
                    if paragraphs_count % 10 == 0:  # 每10个段落记录一次进度
                        logger.debug(f"      已处理 {paragraphs_count} 个段落 ({processed_lines}/{total_lines} 行)")
                    
                    para = doc.add_paragraph()
                    
                    # 处理Markdown加粗（**文本**）
                    import re
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    has_bold = any(part.startswith('**') and part.endswith('**') for part in parts)
                    if has_bold:
                        logger.debug(f"      第{line_num}行: 包含Markdown加粗文本")
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            # 加粗文本
                            bold_text = part.strip('*')
                            run = para.add_run(bold_text)
                            run.font.name = default_font_name
                            run.font.size = default_font_size
                            run.font.bold = True
                        elif part.strip():
                            # 普通文本
                            run = para.add_run(part)
                            run.font.name = default_font_name
                            run.font.size = default_font_size
                    
                    # 应用特殊格式（如果是特殊章节）
                    if is_special_section and special_type in special_formats:
                        special_format = special_formats[special_type]
                        
                        # 检查是否是标题行（第一行且较短，可能是标题）
                        if is_first_line and len(line) < 50:
                            # 应用标题格式
                            if 'title_format' in special_format:
                                title_format = special_format['title_format']
                                for run in para.runs:
                                    run.font.name = title_format.get('font_name', '黑体')
                                    run.font.size = Pt(title_format.get('font_size', 15))
                                    run.font.bold = title_format.get('bold', True)
                                
                                alignment = title_format.get('alignment', 'center')
                                if alignment == 'center':
                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                elif alignment == 'right':
                                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                else:
                                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                                # 如果是目录标题，应用标题后的空行规则
                                if special_type == 'table_of_contents' and toc_title_after_lines > 0:
                                    for _ in range(toc_title_after_lines):
                                        doc.add_paragraph()
                                    logger.info(f"      第{line_num}行后: 目录标题后添加 {toc_title_after_lines} 个空行")
                            is_first_line = False
                        else:
                            # 应用内容格式
                            if 'content_format' in special_format:
                                content_format = special_format['content_format']
                                for run in para.runs:
                                    run.font.name = content_format.get('font_name', '宋体')
                                    run.font.size = Pt(content_format.get('font_size', 12))
                                
                                if 'line_spacing' in content_format:
                                    para.paragraph_format.line_spacing = content_format['line_spacing']
                            elif 'font_name' in special_format:  # 关键词等简单格式
                                for run in para.runs:
                                    run.font.name = special_format.get('font_name', '宋体')
                                    run.font.size = Pt(special_format.get('font_size', 12))
                                
                                if 'line_spacing' in special_format:
                                    para.paragraph_format.line_spacing = special_format['line_spacing']
                                
                                # 目录条目格式应用（支持多级目录格式）
                                if special_type == 'table_of_contents':
                                    # 检查是否有level_formats（多级目录格式）
                                    level_formats = special_format.get('level_formats', {})
                                    
                                    # 判断当前行的目录级别（通过内容特征）
                                    # 一级目录：如"一、绪论"、"1 绪论"、"一 、绪论"（注意空格）
                                    # 二级目录：如"（一）"、"1.1"、"（一）"
                                    # 三级目录：如"1.1.1"
                                    import re
                                    line_text = line.strip()
                                    toc_level = 1  # 默认一级
                                    
                                    # 判断目录级别（更精确的匹配）
                                    if re.match(r'^[一二三四五六七八九十]+[、．\s]', line_text) or re.match(r'^\d+[、．\s]', line_text):
                                        # 一级目录（中文编号或数字编号开头，可能有空格）
                                        toc_level = 1
                                    elif re.match(r'^[（(][一二三四五六七八九十]+[）)]', line_text) or re.match(r'^\d+\.\d+', line_text):
                                        # 二级目录（（一）或1.1格式）
                                        toc_level = 2
                                    elif re.match(r'^\d+\.\d+\.\d+', line_text):
                                        # 三级目录（1.1.1格式）
                                        toc_level = 3
                                    
                                    # 应用对应级别的目录格式（优先级：level_formats > entry_format）
                                    level_key = f'level{toc_level}'
                                    applied_format = None
                                    
                                if level_formats and level_key in level_formats:
                                    # 使用对应级别的格式
                                    applied_format = level_formats[level_key]
                                    logger.debug(f"      第{line_num}行: 应用目录格式 - 级别{toc_level}, 字体={applied_format.get('font_name', 'N/A')}, 字号={applied_format.get('font_size', 'N/A')}磅")
                                elif 'entry_format' in special_format:
                                    # 使用通用目录条目格式
                                    applied_format = special_format['entry_format']
                                    logger.debug(f"      第{line_num}行: 应用目录通用格式 - 字体={applied_format.get('font_name', 'N/A')}, 字号={applied_format.get('font_size', 'N/A')}磅")
                                    
                                    if applied_format:
                                        # 应用字体格式
                                        for run in para.runs:
                                            if 'font_name' in applied_format:
                                                run.font.name = applied_format['font_name']
                                            if 'font_size' in applied_format:
                                                run.font.size = Pt(applied_format['font_size'])
                                        
                                        # 应用缩进（level_formats中有indent字段）
                                        if 'indent' in applied_format:
                                            para.paragraph_format.left_indent = Pt(applied_format['indent'])
                                        
                                        # 应用行距
                                        if 'line_spacing' in applied_format:
                                            para.paragraph_format.line_spacing = applied_format['line_spacing']
                                        
                                        # 应用对齐方式
                                        alignment = applied_format.get('alignment', 'justify')
                                        if alignment == 'justify':
                                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                        elif alignment == 'center':
                                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        elif alignment == 'right':
                                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                        else:
                                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        
                                        # 应用段间距
                                        if 'spacing_before' in applied_format:
                                            para.paragraph_format.space_before = Pt(applied_format['spacing_before'])
                                        if 'spacing_after' in applied_format:
                                            para.paragraph_format.space_after = Pt(applied_format['spacing_after'])
                                is_first_line = False
                    else:
                        # 应用普通段落格式
                        alignment = para_config.get('alignment', 'left')
                        if alignment == 'center':
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif alignment == 'right':
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif alignment == 'justify':
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # 设置行距
                        if 'line_spacing' in para_config:
                            line_spacing_value = para_config['line_spacing']
                            if isinstance(line_spacing_value, (int, float)):
                                para.paragraph_format.line_spacing = line_spacing_value
                            elif isinstance(line_spacing_value, str):
                                try:
                                    para.paragraph_format.line_spacing = float(line_spacing_value)
                                except ValueError:
                                    para.paragraph_format.line_spacing = 1.5
                        
                        # 设置段间距
                        if 'spacing_before' in para_config:
                            para.paragraph_format.space_before = Pt(para_config['spacing_before'])
                        if 'spacing_after' in para_config:
                            para.paragraph_format.space_after = Pt(para_config['spacing_after'])
                    
                    # 设置首行缩进
                    # 检查是否是首段，以及是否需要首行缩进
                    is_first_para = (line == content_lines[0].strip() if content_lines else False)
                    first_para_indent = paragraph_spacing.get('first_paragraph_indent', True)
                    
                    if 'first_line_indent' in para_config:
                        # 如果首段不需要缩进，且这是首段，则不缩进
                        if is_first_para and not first_para_indent:
                            para.paragraph_format.first_line_indent = Pt(0)
                        else:
                            para.paragraph_format.first_line_indent = Pt(para_config['first_line_indent'])
                    
                    # 段落之间的空行处理
                    if not is_first_para:
                        between_paragraphs = paragraph_spacing.get('between_paragraphs', 0)
                        if between_paragraphs > 0:
                            # 在段落之间添加空行（已在循环外通过doc.add_paragraph()处理）
                            pass
                
                # 应用特殊章节后的布局规则
                if is_special_section and special_type in section_spacing:
                    section_spacing_config = section_spacing[special_type]
                    after_section = section_spacing_config.get('after', 0)
                    for _ in range(after_section):
                        doc.add_paragraph()
                    if after_section > 0:
                        logger.debug(f"    特殊章节后添加 {after_section} 个空行")
                
                # 章节结束后的空行
                after_chapter = chapter_spacing.get('after_chapter', 0)
                for _ in range(after_chapter):
                    doc.add_paragraph()
                if after_chapter > 0:
                    logger.debug(f"    章节后添加 {after_chapter} 个空行")
                
                logger.info(f"    章节 {idx+1} 处理完成: 处理了 {processed_lines} 行内容, {paragraphs_count} 个段落, {markdown_headings_count} 个Markdown标题")
        
            # 保存文档
            logger.info(f"[步骤6/6] 保存格式化文档")
            output_dir = Path('uploads/thesis/formatted')
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f'thesis_{thesis_id}_formatted.docx'
            
            # 确保路径是绝对路径
            output_path = output_path.resolve()
            logger.info(f"  输出路径: {output_path}")
            
            # 使用UTF-8编码保存路径，避免编码问题
            try:
                logger.info(f"  开始保存文档到磁盘...")
                doc.save(str(output_path))
                
                # 检查文件是否成功保存
                if output_path.exists():
                    file_size = output_path.stat().st_size
                    logger.info(f"  文档保存成功: {output_path}, 文件大小: {file_size / 1024:.2f} KB")
                else:
                    logger.error(f"  文档保存失败: 文件不存在")
                    raise ServiceException(message='文档保存失败: 文件不存在')
            except Exception as save_error:
                logger.error(f"  保存格式化文档失败: {str(save_error)}", exc_info=True)
                # 如果保存失败，删除可能已创建的不完整文件
                if output_path.exists():
                    try:
                        output_path.unlink()
                        logger.warning(f"  已删除不完整的文件: {output_path}")
                    except:
                        pass
                raise ServiceException(message=f'保存格式化文档失败: {str(save_error)}')
            
            logger.info(f"[格式化完成] 论文ID: {thesis_id}, 输出文件: {output_path}")
            return str(output_path)
            
        except ServiceException:
            # 重新抛出ServiceException，不包装
            raise
        except Exception as e:
            # 捕获所有其他异常，记录详细错误信息
            logger.error(f"创建格式化文档时发生错误: {str(e)}", exc_info=True)
            raise ServiceException(message=f'创建格式化文档失败: {str(e)}')
    
    @classmethod
    def _generate_table_of_contents(
        cls,
        chapters: list,
        format_config: Dict[str, Any],
        layout_rules: Dict[str, Any] = None
    ) -> Any:
        """
        自动生成目录章节
        
        :param chapters: 章节列表（已完成状态）
        :param format_config: 格式配置
        :param layout_rules: 布局规则
        :return: 目录章节对象（包含标题和条目列表），如果生成失败返回None
        """
        from utils.log_util import logger
        
        logger.info("[目录生成] 开始自动生成目录")
        
        # 1. 从格式配置中读取目录格式
        special_sections = format_config.get('special_sections', {})
        toc_config = special_sections.get('table_of_contents', {})
        
        # 获取目录标题文本
        toc_title = toc_config.get('title_text', '目 录')
        
        # 获取目录生成规则
        application_rules = format_config.get('application_rules', {})
        toc_rules = application_rules.get('toc_generation_rules', {})
        include_levels = toc_rules.get('include_levels', [1, 2, 3])
        exclude_sections = toc_rules.get('exclude_sections', ['摘要', '关键词', '目录'])
        
        # 2. 扫描章节，提取标题和级别
        toc_entries = []
        for chapter in chapters:
            # 跳过排除的章节
            chapter_title = getattr(chapter, 'title', '')
            if any(exclude in chapter_title for exclude in exclude_sections):
                logger.debug(f"  跳过排除章节: {chapter_title}")
                continue
            
            # 获取章节级别
            chapter_level = getattr(chapter, 'level', 1)
            
            # 只包含指定级别的章节
            if chapter_level not in include_levels:
                logger.debug(f"  跳过级别 {chapter_level} 的章节: {chapter_title} (不在包含级别中)")
                continue
            
            # 创建目录条目
            entry = {
                'title': chapter_title,
                'level': chapter_level,
                'page_number': None  # 页码需要在文档生成后确定
            }
            toc_entries.append(entry)
            logger.debug(f"  添加目录条目: {chapter_title} (级别: {chapter_level})")
        
        if len(toc_entries) == 0:
            logger.warning("[目录生成] 没有可用的目录条目，跳过目录生成")
            return None
        
        logger.info(f"[目录生成] 共生成 {len(toc_entries)} 个目录条目")
        
        # 3. 创建目录章节对象（模拟章节对象）
        class TOCChapter:
            def __init__(self, title, entries):
                self.title = title
                self.content = ''  # 目录内容将在格式化时生成
                self.level = 0  # 目录是特殊章节，级别为0
                self.order_num = -1  # 目录应该在第一个位置
                self.toc_entries = entries  # 目录条目列表
                self.is_toc = True  # 标记为目录章节
        
        return TOCChapter(toc_title, toc_entries)
    
    @classmethod
    def _number_to_chinese(cls, num: int) -> str:
        """
        将数字转换为中文（支持1-99）
        
        :param num: 数字
        :return: 中文数字
        """
        chinese_nums = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        
        if num <= 0:
            return ''
        elif num <= 10:
            return chinese_nums[num]
        elif num < 20:
            return '十' + (chinese_nums[num % 10] if num % 10 > 0 else '')
        elif num < 100:
            tens = num // 10
            ones = num % 10
            if ones == 0:
                return chinese_nums[tens] + '十'
            else:
                return chinese_nums[tens] + '十' + chinese_nums[ones]
        else:
            # 超过99，返回数字本身
            return str(num)
    
    @classmethod
    def _convert_chapter_numbering(
        cls,
        chapters: list,
        conversion_rules: Dict[str, Any]
    ) -> list:
        """
        转换章节编号格式（如：1.1 → 第一章）
        
        :param chapters: 章节列表
        :param conversion_rules: 转换规则
        :return: 转换后的章节列表
        """
        import re
        from utils.log_util import logger
        
        if not conversion_rules.get('enabled', False):
            return chapters
        
        logger.info("[章节编号转换] 开始转换章节编号")
        
        converted_chapters = []
        special_chapters_config = conversion_rules.get('special_chapters', {})
        
        for chapter in chapters:
            title = getattr(chapter, 'title', '')
            original_title = title
            
            # 匹配 X.Y 格式的标题（如：1.1, 2.1）
            pattern = conversion_rules.get('conversion_pattern', r'^(\d+)\.\d+\s+(.+)$')
            match = re.match(pattern, title)
            
            if match:
                # 提取章节编号和标题文本
                source_num = int(match.group(1))
                title_text = match.group(2).strip()
                
                # 转换为"第X章"格式
                chinese_num = cls._number_to_chinese(source_num)
                chapter.title = f"第{chinese_num}章 {title_text}"
                chapter.level = 1  # 设置为一级标题
                
                logger.debug(f"  转换: {original_title} → {chapter.title}")
            else:
                # 检查是否是特殊章节（结论、参考文献、致谢）
                is_special = False
                
                # 检查结论/结语
                conclusion_config = special_chapters_config.get('conclusion', {})
                if conclusion_config.get('remove_numbering', False):
                    source_names = conclusion_config.get('source_names', [])
                    target_name = conclusion_config.get('target_name', '结语')
                    for source_name in source_names:
                        if source_name in title:
                            # 移除章节编号
                            title = re.sub(r'^第[一二三四五六七八九十]+章\s+', '', title)
                            title = re.sub(r'^\d+\.\d+\s+', '', title)
                            # 替换为目标名称
                            if source_name != target_name:
                                title = title.replace(source_name, target_name)
                            chapter.title = title
                            chapter.level = 0  # 特殊章节，无层级
                            is_special = True
                            logger.debug(f"  特殊章节转换: {original_title} → {chapter.title}")
                            break
                
                # 检查参考文献
                if not is_special:
                    references_config = special_chapters_config.get('references', {})
                    if references_config.get('remove_numbering', False):
                        source_names = references_config.get('source_names', [])
                        for source_name in source_names:
                            if source_name in title:
                                # 移除章节编号
                                title = re.sub(r'^第[一二三四五六七八九十]+章\s+', '', title)
                                title = re.sub(r'^\d+\.\d+\s+', '', title)
                                chapter.title = title
                                chapter.level = 0  # 特殊章节，无层级
                                is_special = True
                                logger.debug(f"  特殊章节转换: {original_title} → {chapter.title}")
                                break
                
                # 检查致谢
                if not is_special:
                    if '致谢' in title or '致　谢' in title:
                        # 移除章节编号
                        title = re.sub(r'^第[一二三四五六七八九十]+章\s+', '', title)
                        title = re.sub(r'^\d+\.\d+\s+', '', title)
                        chapter.title = title
                        chapter.level = 0  # 特殊章节，无层级
                        logger.debug(f"  特殊章节转换: {original_title} → {chapter.title}")
            
            converted_chapters.append(chapter)
        
        logger.info(f"[章节编号转换] 转换完成，共 {len(converted_chapters)} 个章节")
        return converted_chapters
    
    @classmethod
    def _extract_abstract_and_keywords(
        cls,
        chapters: list,
        format_config: Dict[str, Any],
        extraction_rules: Dict[str, Any]
    ) -> tuple[list, Dict[str, Any]]:
        """
        从第一章中提取摘要和关键词，创建独立章节
        
        :param chapters: 章节列表
        :param format_config: 格式配置
        :param extraction_rules: 提取规则
        :return: (更新后的章节列表, 摘要和关键词信息)
        """
        import re
        from utils.log_util import logger
        
        if not extraction_rules.get('enabled', False) or not chapters:
            return chapters, {}
        
        logger.info("[摘要提取] 开始从第一章提取摘要和关键词")
        
        first_chapter = chapters[0]
        content = getattr(first_chapter, 'content', '')
        title = getattr(first_chapter, 'title', '')
        
        # 检查第一章是否已经包含摘要内容
        # 如果第一章标题已经是"摘要"或"Abstract"，则不需要提取
        if '摘要' in title or 'Abstract' in title:
            logger.info("  第一章已经是摘要章节，跳过提取")
            return chapters, {}
        
        abstract_info = {}
        keywords_info = {}
        
        # 使用提取模式识别摘要内容
        extraction_pattern = extraction_rules.get('extraction_pattern', r'(摘要|Abstract)[：:]\s*(.+?)(?=关键词|Key words|$)')
        abstract_match = re.search(extraction_pattern, content, re.DOTALL | re.IGNORECASE)
        
        if abstract_match:
            abstract_text = abstract_match.group(2).strip()
            # 清理文本（移除多余的空白字符）
            abstract_text = re.sub(r'\s+', ' ', abstract_text)
            abstract_text = abstract_text.strip()
            
            if abstract_text:
                abstract_info = {
                    'title': '摘要',
                    'content': abstract_text,
                    'level': 0,
                    'order_num': -2  # 在目录之后
                }
                logger.info(f"  提取到摘要内容，长度: {len(abstract_text)} 字符")
        
        # 识别关键词
        keywords_pattern = r'(关键词|Key words)[：:]\s*(.+?)(?=\n\n|\n\s*\n|$)'
        keywords_match = re.search(keywords_pattern, content, re.DOTALL | re.IGNORECASE)
        
        if keywords_match:
            keywords_text = keywords_match.group(2).strip()
            # 清理文本
            keywords_text = re.sub(r'\s+', ' ', keywords_text)
            keywords_text = keywords_text.strip()
            
            if keywords_text:
                keywords_info = {
                    'title': '关键词',
                    'content': keywords_text,
                    'level': 0,
                    'order_num': -1  # 在摘要之后
                }
                logger.info(f"  提取到关键词内容: {keywords_text[:50]}...")
        
        # 从第一章中移除摘要和关键词内容
        if abstract_match or keywords_match:
            # 移除摘要部分
            if abstract_match:
                content = content[:abstract_match.start()] + content[abstract_match.end():]
            
            # 移除关键词部分
            if keywords_match:
                content = content[:keywords_match.start()] + content[keywords_match.end():]
            
            # 清理内容（移除多余的空白行）
            content = re.sub(r'\n{3,}', '\n\n', content)
            content = content.strip()
            
            # 更新第一章内容
            first_chapter.content = content
            logger.info(f"  已从第一章移除摘要和关键词内容，剩余内容长度: {len(content)} 字符")
        
        return chapters, {'abstract': abstract_info, 'keywords': keywords_info}
    
    @classmethod
    def _create_chapter_from_info(cls, chapter_info: Dict[str, Any]) -> Any:
        """
        根据章节信息创建章节对象
        
        :param chapter_info: 章节信息字典
        :return: 章节对象
        """
        class Chapter:
            def __init__(self, info):
                self.title = info.get('title', '')
                self.content = info.get('content', '')
                self.level = info.get('level', 0)
                self.order_num = info.get('order_num', 0)
                self.status = 'completed'
        
        return Chapter(chapter_info)